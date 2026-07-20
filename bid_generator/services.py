from io import BytesIO
import re
from html import escape

from django.core.files.base import ContentFile
from django.http import HttpResponse
from django.utils import timezone
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import PageBreak, Paragraph, Spacer, Table, TableStyle

from core.models import SystemSettings
from core.pdf import add_signature, apply_company_letterhead_pdf, build_pdf_response, letterhead_elements, styles
from documents.models import CompanyDocument
from tenders.services import bid_pack_output_for_task, required_document_types
from .models import BidDocument


ATTACHMENT_ORDER = [
    CompanyDocument.DocumentType.PACRA,
    CompanyDocument.DocumentType.TPIN_CERTIFICATE,
    CompanyDocument.DocumentType.ZRA_TAX_CLEARANCE,
    CompanyDocument.DocumentType.NAPSA,
    CompanyDocument.DocumentType.WORKERS_COMPENSATION,
    CompanyDocument.DocumentType.NCC,
    CompanyDocument.DocumentType.NCC_B,
    CompanyDocument.DocumentType.NCC_R,
    CompanyDocument.DocumentType.NCC_E,
    CompanyDocument.DocumentType.ERB,
    CompanyDocument.DocumentType.EIZ_CERTIFICATE,
    CompanyDocument.DocumentType.ZEMA_LICENSE,
    CompanyDocument.DocumentType.ROADWORTHINESS,
    CompanyDocument.DocumentType.ZPPA_REGISTRATION,
    CompanyDocument.DocumentType.BANK_CONFIRMATION,
    CompanyDocument.DocumentType.AUDITED_FINANCIALS,
    CompanyDocument.DocumentType.DELIVERY_EVIDENCE,
    CompanyDocument.DocumentType.TRAINING_PROGRAMME,
    CompanyDocument.DocumentType.WARRANTY_UNDERTAKING,
    CompanyDocument.DocumentType.PAST_CONTRACT,
    CompanyDocument.DocumentType.COMPANY_PROFILE,
    CompanyDocument.DocumentType.OTHER,
]

QUALIFICATION_FORM_DEFINITIONS = {
    'ELI 1.1': {
        'title': 'Form ELI 1.1 - Bidder Information Sheet',
        'factor': 'Eligibility',
        'headers': ['Field', 'TenderAI filled value'],
        'rows': [
            ('Bidder legal name', 'company.name'),
            ('Country of registration', 'Zambia'),
            ('Year of registration', 'To be completed from PACRA certificate'),
            ('Legal address', 'company.address'),
            ('Authorized representative', 'To be completed'),
            ('TPIN', 'company.tpin'),
            ('Attached evidence', 'PACRA, TPIN, tax clearance, and other eligibility documents'),
        ],
    },
    'ELI 1.2': {
        'title': 'Form ELI 1.2 - Party to JV Information Sheet',
        'factor': 'Eligibility',
        'headers': ['Field', 'TenderAI filled value'],
        'rows': [
            ('JV party legal name', 'Not applicable unless bidding as a JV'),
            ('Country of registration', 'To be completed if JV applies'),
            ('Legal address', 'To be completed if JV applies'),
            ('Authorized representative', 'To be completed if JV applies'),
        ],
    },
    'CON-2': {
        'title': 'Form CON-2 - Historical Contract Non-Performance',
        'factor': 'Historical Contract Non-Performance',
        'headers': ['Year', 'Contract', 'Non-performance details', 'Status'],
        'rows': [('To be completed', 'To be completed', 'None / disclose if applicable', 'To be completed')],
    },
    'CON-3': {
        'title': 'Form CON-3 - Current Contract Commitments / Works in Progress',
        'factor': 'Historical Contract Non-Performance',
        'headers': ['Client', 'Contract description', 'Value', 'Completion date', 'Percent complete'],
        'rows': [('To be completed', 'Current contract / works in progress', '-', '-', '-')],
    },
    'CCC': {
        'title': 'Form CCC - Current Contract Commitments',
        'factor': 'Historical Contract Non-Performance',
        'headers': ['Client', 'Contract', 'Outstanding value', 'Expected completion'],
        'rows': [('To be completed', 'To be completed', '-', '-')],
    },
    'FIN-3.1': {
        'title': 'Form FIN-3.1 - Financial Situation',
        'factor': 'Financial Situation',
        'headers': ['Financial year', 'Total assets', 'Total liabilities', 'Net worth', 'Source document'],
        'rows': [('To be completed', '-', '-', '-', 'Audited accounts / financial statements')],
    },
    'FIN-3.2': {
        'title': 'Form FIN-3.2 - Average Annual Turnover',
        'factor': 'Financial Situation',
        'headers': ['Year', 'Turnover', 'Currency', 'Source document'],
        'rows': [('To be completed', '-', 'ZMW', 'Financial statements')],
    },
    'FIN-3.3': {
        'title': 'Form FIN-3.3 - Financial Resources',
        'factor': 'Financial Situation',
        'headers': ['Source of financing', 'Amount', 'Currency', 'Evidence'],
        'rows': [('Bank / credit facility / cash resources', '-', 'ZMW', 'Attach bank confirmation or proof')],
    },
    'EXP-2.4.1': {
        'title': 'Form EXP-2.4.1 - General Experience',
        'factor': 'Experience',
        'headers': ['Starting month/year', 'Ending month/year', 'Client', 'Contract description', 'Role'],
        'rows': [('To be completed', 'To be completed', 'To be completed', 'Relevant works/supplies/services', 'Contractor')],
    },
    'EXP-2.4.2(a)': {
        'title': 'Form EXP-2.4.2(a) - Specific Experience',
        'factor': 'Experience',
        'headers': ['Contract', 'Client', 'Country', 'Contract value', 'Completion date', 'Similarity'],
        'rows': [('To be completed', 'To be completed', 'Zambia', '-', '-', 'Similar assignment')],
    },
    'EXP-2.4.2(b)': {
        'title': 'Form EXP-2.4.2(b) - Specific Experience in Key Activities',
        'factor': 'Experience',
        'headers': ['Key activity', 'Contract reference', 'Quantity / scope', 'Completion date'],
        'rows': [('To be completed', 'To be completed', '-', '-')],
    },
    'PER-1': {
        'title': 'Form PER-1 - Proposed Personnel',
        'factor': 'Personnel',
        'headers': ['Position', 'Name', 'Qualification', 'Years experience'],
        'rows': [('Project Manager / Supervisor', 'To be completed', 'To be completed', '-')],
    },
    'PER-2': {
        'title': 'Form PER-2 - Resume of Proposed Personnel',
        'factor': 'Personnel',
        'headers': ['Field', 'TenderAI filled value'],
        'rows': [
            ('Position', 'To be completed'),
            ('Name', 'To be completed'),
            ('Date of birth', 'To be completed'),
            ('Professional qualifications', 'To be completed'),
            ('Present employer', 'To be completed'),
            ('Relevant experience', 'To be completed'),
        ],
    },
    'EQU': {
        'title': 'Forms for Equipment',
        'factor': 'Equipment',
        'headers': ['Equipment type', 'Make / model', 'Capacity', 'Owned / leased', 'Availability'],
        'rows': [('To be completed', '-', '-', '-', 'Available for contract')],
    },
    'MFR': {
        'title': 'Manufacturer’s Authorisation',
        'factor': 'Manufacturer Authorisation',
        'headers': ['Field', 'TenderAI filled value'],
        'rows': [
            ('Manufacturer / supplier', 'To be completed where goods require authorisation'),
            ('Tender title', 'tender.title'),
            ('Bidder name', 'company.name'),
            ('Authorised goods', 'To be completed'),
            ('Signature', 'To be completed'),
        ],
    },
}


def build_bid_checklist(bid_pack):
    tender = bid_pack.tender
    itb_rows = ordered_itb_rows(bid_pack)
    if itb_rows:
        checklist = [required_item for _, required_item, _, _ in itb_rows]
    else:
        checklist = [
            'Cover letter',
            'Company profile summary',
            'Price schedule',
            'Priced schedule / bill of quantities',
        ]

    for requirement in tender.requirements.all():
        if not requirement.is_mandatory:
            continue
        if requirement.requirement_type not in {
            requirement.RequirementType.MANDATORY_DOCUMENT,
            requirement.RequirementType.CERTIFICATE,
            requirement.RequirementType.REQUIRED_FORM,
            requirement.RequirementType.BID_SECURITY,
            requirement.RequirementType.SITE_VISIT,
        }:
            continue
        item = str(requirement.description).strip()
        if not is_clean_checklist_item(item):
            continue
        if itb_rows and requirement.requirement_type == requirement.RequirementType.REQUIRED_FORM:
            continue
        if itb_rows and 'bid security' in item.lower() and any('bid security' in existing.lower() for existing in checklist):
            continue
        if item and item not in checklist:
            checklist.append(item)
    return checklist


def is_clean_checklist_item(item):
    if not item or len(item) < 3:
        return False
    noisy_markers = ['........', 'error! bookmark', 'section vi.', 'section iii.', 'this section contains']
    lower = item.lower()
    if any(marker in lower for marker in noisy_markers):
        return False
    if lower.startswith('('):
        return False
    if lower.startswith(('information that ', 'evaluation of bids')):
        return False
    if lower in {'and schedules', 'schedules'}:
        return False
    if item.count('.') > 8:
        return False
    return True


def requirements_matrix_rows(bid_pack):
    rows = []
    for requirement in bid_pack.tender.requirements.all():
        if requirement.requirement_type not in {
            requirement.RequirementType.MANDATORY_DOCUMENT,
            requirement.RequirementType.CERTIFICATE,
            requirement.RequirementType.REQUIRED_FORM,
            requirement.RequirementType.BID_SECURITY,
            requirement.RequirementType.SITE_VISIT,
            requirement.RequirementType.EXPERIENCE,
        }:
            continue
        if not is_clean_checklist_item(requirement.description):
            continue
        rows.append((
            requirement.get_requirement_type_display(),
            requirement.description,
            'Yes' if requirement.is_mandatory else 'No',
            'To confirm / attach',
        ))
    if not rows:
        rows.append((
            'General',
            'No solicitation requirements have been extracted yet. Upload/analyse the solicitation document before final submission.',
            'Yes',
            'Pending analysis',
        ))
    return rows


def ordered_itb_rows(bid_pack):
    rows = []
    tasks = bid_pack.tender.bid_tasks.filter(title__startswith='ITB ').order_by('sort_order', 'title')
    for task in tasks:
        reference, _, required_item = task.title.partition(':')
        rows.append((
            reference.strip(),
            required_item.strip() or task.title,
            bid_pack_output_for_task(task),
            task.get_status_display(),
        ))
    return rows


def xml_bid_structure_groups(bid_pack):
    groups = []
    group_map = {}
    for item in expanded_tender_xml_items(bid_pack.tender):
        envelope = item.get('envelope') or 'Bid requirements'
        if envelope not in group_map:
            group_map[envelope] = {
                'title': envelope,
                'items': [],
            }
            groups.append(group_map[envelope])
        group_map[envelope]['items'].append(item)
    return groups


def expanded_tender_xml_items(tender):
    return expand_xml_bid_items(tender.itb_11_items or [])


def expand_xml_bid_items(items):
    expanded = []
    next_order = 1
    for item in items or []:
        response_items = xml_response_lines(item)
        if should_split_xml_item(item, response_items):
            for response in response_items:
                expanded.append({
                    **item,
                    'order': next_order,
                    'requirement': clean_split_requirement_label(response),
                    'response': response,
                    'response_items': [response],
                    'title': f'{clean_split_requirement_label(response)} - {response}',
                    'parent_reference': item.get('reference') or '',
                    'reference': item.get('reference') or f'XML {next_order}',
                })
                next_order += 1
            continue
        if len(response_items) == 1 and is_generic_xml_requirement(item):
            response = response_items[0]
            expanded.append({
                **item,
                'order': next_order,
                'requirement': clean_split_requirement_label(response),
                'response': response,
                'response_items': [response],
                'title': f'{clean_split_requirement_label(response)} - {response}',
                'parent_reference': item.get('reference') or '',
                'reference': item.get('reference') or f'XML {next_order}',
            })
            next_order += 1
            continue
        expanded.append({**item, 'order': next_order})
        next_order += 1
    return expanded


def should_split_xml_item(item, response_items):
    if len(response_items) <= 1:
        return False
    return is_generic_xml_requirement(item)


def is_generic_xml_requirement(item):
    requirement = normalize_match_text(item.get('requirement') or item.get('title') or '')
    generic_words = [
        'preliminary requirements',
        'preliminiary requirements',
        'preliminery requirements',
        'preliminary evaluation',
        'preliminiary evaluation',
        'preliminery evaluation',
        'preliminary examination',
        'preliminiary examination',
        'preliminery examination',
        'evaluation criteria',
        'examination criteria',
        'mandatory requirements',
        'proof documents',
        'technical requirements',
        'technical evaluation',
        'technical examination',
        'commercial requirements',
        'commercial evaluation',
        'commercial examination',
        'commercial offer',
        'technical response',
        'requirements',
        'criteria',
    ]
    return any(word in requirement for word in generic_words)


def clean_split_requirement_label(response):
    text = re.sub(r'\s+', ' ', str(response or '')).strip(' .;:-')
    normal = normalize_match_text(text)
    if 'signed quotation' in normal or 'submit a quotation' in normal:
        return 'Signed Quotation'
    if 'power of attorney' in normal:
        return 'Power of Attorney'
    if 'bid submission form' in normal or 'bid form' in normal:
        return 'Bid Submission Form'
    if 'price schedule' in normal:
        return 'Price Schedule'
    if 'pacra printout' in normal or 'beneficial ownership' in normal:
        return 'PACRA Printout / Beneficial Ownership'
    if 'pacra certificate' in normal or 'certificate of incorporation' in normal:
        return 'PACRA Certificate of Incorporation'
    if 'zra tax' in normal or 'tax clearance' in normal or 'zra clearance' in normal:
        return 'ZRA Tax Clearance Certificate'
    if 'napsa' in normal:
        return 'NAPSA Compliance Certificate'
    if 'workers compensation' in normal:
        return 'Workers Compensation Compliance Certificate'
    if 'technical specification' in normal or 'technical specifications' in normal:
        return 'Technical Specifications Response'
    if 'delivery period' in normal:
        return 'Delivery Period Confirmation'
    if 'payment' in normal:
        return 'Payment Terms Confirmation'
    for _ in range(3):
        text = re.sub(
            r'^(bidders?\s+(?:to|shall|must|should)\s+|submit|attach|provide|state|indicate|copy of|bidders? must provide|bidders? should submit)\s+',
            '',
            text,
            flags=re.I,
        )
        text = re.sub(r'^(duly\s+)?signed\s+', '', text, flags=re.I)
    text = re.split(r'\s+-\s+', text, maxsplit=1)[0].strip(' .;:-')
    if len(text) > 130:
        text = text[:127].rsplit(' ', 1)[0] + '...'
    return text or 'Bid requirement'


def has_xml_bid_structure(bid_pack):
    return bool(xml_bid_structure_groups(bid_pack))


def table_of_contents(bid_pack=None):
    has_xml = bool(bid_pack and has_xml_bid_structure(bid_pack))
    sections = [
        'Tender and Bidder Information',
        'Submission Readiness Gaps',
        'Documents Comprising the Bid',
        'Solicitation Forms and Statements',
        'Proof Documents / Eligibility Evidence',
        'Qualification Schedules',
        'Technical Response',
        'Commercial Offer / Price Schedule',
        'Company Profile Summary',
        'Final Signature Section',
    ]
    if has_xml:
        sections.insert(4, 'XML Envelope Submission Order')
    if bid_pack and ordered_certificate_documents(bid_pack):
        sections.append('Attached Company Certificates')
    return sections


def submission_order_rows(bid_pack):
    rows = []
    if has_xml_bid_structure(bid_pack):
        for group in xml_bid_structure_groups(bid_pack):
            for item in group['items']:
                response_lines = xml_response_lines(item)
                rows.append({
                    'order': item.get('order') or len(rows) + 1,
                    'envelope': group['title'],
                    'requirement': item.get('requirement') or item.get('title') or 'Bid requirement',
                    'response': '; '.join(response_lines[:3]) if response_lines else 'Complete / attach as required',
                    'prepared_as': prepared_output_for_xml_item(bid_pack, item),
                    'mandatory': bool(item.get('mandatory')),
                })
        return rows

    for reference, required_item, output, status in ordered_itb_rows(bid_pack):
        rows.append({
            'order': reference,
            'envelope': 'ITB checklist',
            'requirement': required_item,
            'response': status,
            'prepared_as': output,
            'mandatory': True,
        })
    return rows


def xml_response_lines(item):
    lines = [line.strip() for line in item.get('response_items') or [] if str(line).strip()]
    if not lines and item.get('response'):
        lines = split_grouped_xml_response_text(item.get('response'))
    elif len(lines) == 1 and ';' in lines[0]:
        lines = split_grouped_xml_response_text(lines[0])
    return lines


def split_grouped_xml_response_text(response):
    response = re.sub(r'\s+', ' ', str(response or '')).strip()
    if not response:
        return []
    parts = re.split(
        r';\s*(?=(?:Bidders?\s+(?:to|shall|must|should)|Attach|Provide|Indicate|Submit|Complete|Signed?|Valid|Availability|Evidence|Documentary|Warranty|Delivery|Payment)\b)',
        response,
        flags=re.I,
    )
    if len(parts) == 1 and ';' in response:
        parts = response.split(';')
    return [part.strip(' ;') for part in parts if part.strip(' ;')][:30]


def prepared_output_for_xml_item(bid_pack, item):
    text = normalize_match_text(' '.join([
        str(item.get('requirement', '')),
        str(item.get('response', '')),
        ' '.join(item.get('response_items') or []),
    ]))
    form_code = qualification_form_code_for_text(text)
    if 'bid declaration' in text or 'bid securing' in text or 'bid security' in text:
        return 'Bid-Securing Declaration form'
    if is_letter_of_bid_requirement(text):
        if is_price_schedule_requirement(text):
            return 'Letter of Bid plus Price Schedule'
        return 'Letter of Bid / Tender Submission Letter'
    if is_price_schedule_requirement(text):
        return 'Signed quotation / price schedule'
    if is_power_of_attorney_requirement(text):
        return 'Power of Attorney / Signatory Authorisation form'
    if form_code:
        if form_code == 'MFR':
            return "Manufacturer's Authorisation form"
        return f'Qualification form {form_code}'
    matched_document = match_company_document_for_xml_item(bid_pack, item)
    if matched_document:
        return f'Attach {matched_document.get_document_type_display()}'
    if any(word in text for word in ['technical', 'specification', 'compliance', 'capacity']):
        return 'Technical response schedule'
    if any(word in text for word in ['price', 'commercial', 'payment', 'warranty', 'delivery']):
        return 'Commercial / undertaking schedule'
    return 'Prepared response form'


def submission_gap_rows(bid_pack):
    return [
        (item['label'], item['status'], item['action'])
        for item in submission_gap_items(bid_pack)
    ]


def submission_gap_items(bid_pack):
    docs = list(bid_pack.company.documents.all())
    required = list(required_document_types(bid_pack.tender))
    if not required:
        required = [
            CompanyDocument.DocumentType.PACRA,
            CompanyDocument.DocumentType.ZRA_TAX_CLEARANCE,
            CompanyDocument.DocumentType.TPIN_CERTIFICATE,
            CompanyDocument.DocumentType.ZPPA_REGISTRATION,
            CompanyDocument.DocumentType.BANK_CONFIRMATION,
            CompanyDocument.DocumentType.COMPANY_PROFILE,
        ]
    items = []
    for doc_type in required:
        matching = [document for document in docs if document.document_type == doc_type]
        label = CompanyDocument.DocumentType(doc_type).label
        if not matching:
            items.append({
                'label': label,
                'status': 'Missing',
                'action': 'Upload this document before final bid submission.',
                'document_type': doc_type,
            })
            continue
        if any(document.is_expired for document in matching):
            items.append({
                'label': label,
                'status': 'Expired',
                'action': 'Renew this document before submission.',
                'document_type': doc_type,
            })
    if not bid_pack.company.tpin:
        items.append({
            'label': 'Company TPIN',
            'status': 'Missing',
            'action': 'Fill in the company TPIN on the company profile.',
            'document_type': '',
        })
    if not bid_pack.company.registration_number:
        items.append({
            'label': 'PACRA registration number',
            'status': 'Missing',
            'action': 'Fill in the company registration number.',
            'document_type': '',
        })
    if not bid_pack.company.address:
        items.append({
            'label': 'Company address',
            'status': 'Missing',
            'action': 'Fill in the company address for letterhead and forms.',
            'document_type': '',
        })
    return items


def company_document_checklist(company):
    today_docs = company.documents.all()
    required = [
        CompanyDocument.DocumentType.PACRA,
        CompanyDocument.DocumentType.ZRA_TAX_CLEARANCE,
        CompanyDocument.DocumentType.TPIN_CERTIFICATE,
        CompanyDocument.DocumentType.NAPSA,
        CompanyDocument.DocumentType.WORKERS_COMPENSATION,
        CompanyDocument.DocumentType.NCC,
        CompanyDocument.DocumentType.NCC_B,
        CompanyDocument.DocumentType.NCC_R,
        CompanyDocument.DocumentType.NCC_E,
        CompanyDocument.DocumentType.ERB,
        CompanyDocument.DocumentType.EIZ_CERTIFICATE,
        CompanyDocument.DocumentType.ZEMA_LICENSE,
        CompanyDocument.DocumentType.ROADWORTHINESS,
        CompanyDocument.DocumentType.ZPPA_REGISTRATION,
        CompanyDocument.DocumentType.BANK_CONFIRMATION,
        CompanyDocument.DocumentType.AUDITED_FINANCIALS,
        CompanyDocument.DocumentType.DELIVERY_EVIDENCE,
        CompanyDocument.DocumentType.TRAINING_PROGRAMME,
        CompanyDocument.DocumentType.WARRANTY_UNDERTAKING,
        CompanyDocument.DocumentType.PAST_CONTRACT,
        CompanyDocument.DocumentType.COMPANY_PROFILE,
    ]
    rows = []
    for doc_type in required:
        docs = [doc for doc in today_docs if doc.document_type == doc_type]
        status = 'Missing'
        if docs:
            status = 'Available'
            if any(doc.is_expired for doc in docs):
                status = 'Expired'
        rows.append((CompanyDocument.DocumentType(doc_type).label, status))
    return rows


def required_form_specs(bid_pack):
    tender = bid_pack.tender
    company = bid_pack.company
    form_text = ' '.join(
        requirement.description.lower()
        for requirement in tender.requirements.filter(requirement_type='REQUIRED_FORM')
    )
    itb_text = ' '.join(row[1].lower() for row in ordered_itb_rows(bid_pack))
    corpus = f'{form_text} {itb_text} {tender.title.lower()}'

    specs = [
        {
            'title': 'Letter of Bid',
            'kind': 'letter_of_bid',
        }
    ]

    if 'bid securing declaration' in corpus or 'bid security' in corpus:
        specs.append({
            'title': 'Bid-Securing Declaration',
            'kind': 'bid_securing_declaration',
        })

    if 'schedule' in corpus or 'bill of quantities' in corpus or 'boq' in corpus:
        specs.append({
            'title': 'Price Schedule / Bill of Quantities',
            'headers': ['Item No.', 'Description', 'Unit', 'Quantity', 'Unit Price', 'Total'],
            'rows': [
                ('1', 'Complete from solicitation BOQ / price schedule', '-', '-', '-', '-'),
                ('2', 'Add additional line items as required', '-', '-', '-', '-'),
            ],
        })

    if 'authorization' in corpus or 'authorizing' in corpus or 'power of attorney' in corpus:
        specs.append({
            'title': 'Signatory Authorization / Power of Attorney',
            'rows': [
                ('Company', company.name),
                ('Tender title', tender.title),
                ('Authorised person', 'To be completed'),
                ('Position', 'To be completed'),
                ('Scope', 'Authority to sign and submit bid documents'),
                ('Director / witness signature', 'To be completed and signed'),
            ],
        })

    if 'experience' in corpus or tender.requirements.filter(description__icontains='Past Contracts').exists():
        specs.append({
            'title': 'Similar Experience Table',
            'headers': ['Client', 'Contract description', 'Year', 'Value', 'Contact person'],
            'rows': [
                ('To be completed', 'Relevant similar contract', '-', '-', '-'),
                ('To be completed', 'Relevant similar contract', '-', '-', '-'),
            ],
        })

    for form in qualification_form_specs(bid_pack):
        if not any(existing.get('title') == form.get('title') for existing in specs):
            specs.append(form)

    unmapped = unmapped_required_items(bid_pack)
    if unmapped:
        specs.append({
            'title': 'Other Required Items Not Yet Mapped',
            'headers': ['Requirement', 'Action'],
            'rows': [(item, 'Confirm where this must be included in the final bid pack.') for item in unmapped],
        })

    return specs


def qualification_form_specs(bid_pack):
    codes = []
    for requirement in bid_pack.tender.requirements.filter(requirement_type='REQUIRED_FORM'):
        text = requirement.description.upper().replace('–', '-')
        for code in QUALIFICATION_FORM_DEFINITIONS:
            if code.upper() in text and code not in codes:
                codes.append(code)
    document = bid_pack.tender.solicitation_documents.order_by('-uploaded_at').first()
    analysis = document.analysis_summary if document else {}
    for form in analysis.get('qualification_forms', []):
        code = form.get('code')
        if code in QUALIFICATION_FORM_DEFINITIONS and code not in codes:
            codes.append(code)
    return [build_qualification_form_spec(bid_pack, code) for code in codes]


def build_qualification_form_spec(bid_pack, code):
    definition = QUALIFICATION_FORM_DEFINITIONS[code]
    return {
        'title': definition['title'],
        'factor': definition['factor'],
        'headers': definition['headers'],
        'rows': [
            tuple(resolve_form_value(bid_pack, value) for value in row)
            for row in definition['rows']
        ],
    }


def resolve_form_value(bid_pack, value):
    if value == 'company.name':
        return bid_pack.company.name
    if value == 'company.address':
        return bid_pack.company.address or '-'
    if value == 'company.tpin':
        return bid_pack.company.tpin or '-'
    if value == 'tender.title':
        return bid_pack.tender.title
    return value


def unmapped_required_items(bid_pack):
    mapped_words = [
        'pacra', 'zra', 'tax', 'napsa', 'workers', 'ncc', 'eiz', 'bank', 'past contract',
        'letter of bid', 'bid securing', 'bid security', 'schedule', 'personnel', 'equipment',
        'eli', 'con-', 'fin-', 'exp-', 'per-', 'manufacturer', 'experience',
    ]
    items = []
    for requirement in bid_pack.tender.requirements.filter(is_mandatory=True):
        text = str(requirement.description).strip()
        lower = text.lower()
        if not is_clean_checklist_item(text):
            continue
        if any(word in lower for word in mapped_words):
            continue
        if text not in items:
            items.append(text)
    return items[:20]


def add_pdf_required_form_tables(elements, bid_pack, style):
    for spec in required_form_specs(bid_pack):
        if spec.get('kind') == 'letter_of_bid':
            elements.append(PageBreak())
        if spec.get('kind') == 'bid_securing_declaration':
            elements.append(PageBreak())
        elements.append(Paragraph(spec['title'], style['Heading3']))
        if spec.get('kind') == 'letter_of_bid':
            add_pdf_letter_of_bid(elements, bid_pack, style)
            elements.append(Spacer(1, 10))
            continue
        if spec.get('kind') == 'bid_securing_declaration':
            add_pdf_bid_securing_declaration(elements, bid_pack, style)
            elements.append(Spacer(1, 10))
            elements.append(PageBreak())
            continue
        if 'headers' in spec:
            rows = [spec['headers'], *spec['rows']]
            col_width = 475 / len(spec['headers'])
            elements.append(Table(rows, colWidths=[col_width] * len(spec['headers']), style=form_table_style()))
        else:
            rows = [['Field', 'TenderAI filled value']]
            rows.extend([[field, Paragraph(str(value), style['BodyText'])] for field, value in spec['rows']])
            elements.append(Table(rows, colWidths=[150, 325], style=form_table_style()))
        elements.append(Spacer(1, 10))


def add_docx_required_form_tables(document, bid_pack):
    for spec in required_form_specs(bid_pack):
        if spec.get('kind') == 'letter_of_bid':
            document.add_page_break()
        if spec.get('kind') == 'bid_securing_declaration':
            document.add_page_break()
        document.add_heading(spec['title'], level=2)
        if spec.get('kind') == 'letter_of_bid':
            add_docx_letter_of_bid(document, bid_pack)
            continue
        if spec.get('kind') == 'bid_securing_declaration':
            add_docx_bid_securing_declaration(document, bid_pack)
            document.add_page_break()
            continue
        if 'headers' in spec:
            table = document.add_table(rows=1, cols=len(spec['headers']))
            table.style = 'Table Grid'
            for idx, heading in enumerate(spec['headers']):
                table.rows[0].cells[idx].text = heading
            for row in spec['rows']:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = str(value)
        else:
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Field'
            table.rows[0].cells[1].text = 'TenderAI filled value'
            for field, value in spec['rows']:
                cells = table.add_row().cells
                cells[0].text = str(field)
                cells[1].text = str(value)


def letter_of_bid_context(bid_pack):
    tender = bid_pack.tender
    company = bid_pack.company
    closing_date = tender.closing_at or tender.closing_date or '-'
    references = tender_reference_numbers(tender)
    return {
        'date': timezone.localdate().strftime('%d/%m/%Y'),
        'bidding_no': references['tender_number'],
        'invitation_no': references['invitation_number'],
        'to': tender.procuring_entity or '-',
        'works': tender.title,
        'bid_price': 'As stated in the attached Price Schedule / Financial Offer',
        'discounts': 'None, unless stated separately in the priced schedule',
        'validity': 'the period stated in the solicitation document',
        'company': company.name,
        'company_address': company.address or '-',
        'company_tpin': company.tpin or '-',
        'representative': 'Authorised Representative',
        'signatory_capacity': 'Authorised Representative',
        'submission_method': tender.submission_method or '-',
        'closing_date': closing_date,
    }


def tender_reference_numbers(tender):
    cover_references = latest_cover_references(tender)
    tender_number = cover_references.get('tender_number') or tender.tender_number or tender.zppa_resource_id or '-'
    invitation_number = (
        cover_references.get('invitation_number')
        or cover_references.get('tender_number')
        or tender.tender_number
        or tender.zppa_resource_id
        or '-'
    )
    return {
        'tender_number': tender_number,
        'invitation_number': invitation_number,
    }


def latest_cover_references(tender):
    references = {}
    for document in tender.solicitation_documents.order_by('-uploaded_at'):
        summary = document.analysis_summary or {}
        document_references = summary.get('cover_references') or {}
        for key in ['tender_number', 'invitation_number']:
            value = str(document_references.get(key) or '').strip()
            if value and not references.get(key):
                references[key] = value
        if references.get('tender_number') and references.get('invitation_number'):
            break
    return references


def letter_of_bid_declarations(context):
    return [
        'We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause 8.',
        f'We offer to execute in conformity with the Bidding Documents the following Works: {context["works"]}.',
        f'The total price of our Bid, excluding any discounts offered below is: {context["bid_price"]}.',
        f'The discounts offered and the methodology for their application are: {context["discounts"]}.',
        f'Our bid shall remain valid for {context["validity"]} from the date fixed for the bid submission deadline and shall remain binding upon us.',
        'If price adjustment provisions apply, the Table(s) of Adjustment Data shall be considered part of this Bid.',
        'If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document.',
        'Our firm, including any subcontractors or suppliers for any part of the Contract, have nationalities from eligible countries.',
        'We, including any subcontractors or suppliers for any part of the contract, do not have any conflict of interest in accordance with ITB 4.3.',
        'We are not participating, as a Bidder or as a subcontractor, in more than one bid in this bidding process, other than alternative offers submitted in accordance with ITB 13.',
        'Our firm, its affiliates or subsidiaries, including any subcontractors or suppliers for any part of the contract, has not been declared ineligible by ZPPA or by a United Nations Security Council decision.',
        'We are not a government owned entity / We are a government owned entity but meet the requirements of ITB 4.5. Delete whichever is not applicable.',
        'We have paid, or will pay, no commissions, gratuities, or fees with respect to the bidding process or execution of the Contract unless disclosed below.',
        'We understand that this bid, together with your written acceptance included in your notification of award, shall constitute a binding contract between us until a formal contract is prepared and executed.',
        'We understand that you are not bound to accept the best-evaluated bid or any other bid that you may receive.',
        f'If awarded the contract, the person named below shall act as Contractor Representative: {context["representative"]}.',
    ]


def add_pdf_letter_of_bid(elements, bid_pack, style):
    context = letter_of_bid_context(bid_pack)
    form_style = zppa_form_styles(style)
    elements.append(Paragraph('LETTER OF BID / BID SUBMISSION FORM', form_style['title']))
    elements.append(Spacer(1, 8))
    for line in letter_of_bid_reference_lines(context):
        elements.append(Paragraph(line, form_style['right']))
    elements.append(Spacer(1, 8))
    for line in letter_of_bid_recipient_lines(context):
        elements.append(Paragraph(line, form_style['body']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph('We, the undersigned, declare that:', form_style['body']))
    for index, declaration in enumerate(letter_of_bid_declarations(context), start=1):
        letter = chr(96 + index)
        elements.append(Paragraph(f'({letter}) {declaration}', form_style['indented']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('Commissions, gratuities, or fees:', form_style['body']))
    commission_rows = [
        ['Name of Recipient', 'Address', 'Reason', 'Amount'],
        ['None', '-', '-', '-'],
    ]
    elements.append(Table(commission_rows, colWidths=[130, 130, 130, 85], style=form_table_style()))
    elements.append(Spacer(1, 12))
    for line in letter_of_bid_signature_lines(context):
        elements.append(Paragraph(line, form_style['body']))


def letter_of_bid_field_lines(context):
    return letter_of_bid_reference_lines(context) + letter_of_bid_recipient_lines(context)


def letter_of_bid_reference_lines(context):
    return [
        f'Date: {context["date"]}',
        f'Bidding No.: {context["bidding_no"]}',
        f'Invitation for Bid No.: {context["invitation_no"]}',
    ]


def letter_of_bid_recipient_lines(context):
    return [
        f'To: {context["to"]}',
    ]


def letter_of_bid_signature_lines(context):
    return [
        'Name: ____________________________',
        f'In the capacity of: {context["signatory_capacity"]}',
        'Signed: ____________________________',
        f'Duly authorised to sign the Bid for and on behalf of: {context["company"]}',
        f'Date: {context["date"]}',
    ]


def add_docx_letter_of_bid(document, bid_pack):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    context = letter_of_bid_context(bid_pack)
    p = document.add_paragraph('LETTER OF BID / BID SUBMISSION FORM')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_paragraph_font(p, size=13, bold=True)
    for line in letter_of_bid_reference_lines(context):
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_docx_paragraph_font(p)
    for line in letter_of_bid_recipient_lines(context):
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_docx_paragraph_font(p)
    p = document.add_paragraph('We, the undersigned, declare that:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_docx_paragraph_font(p)
    for index, declaration in enumerate(letter_of_bid_declarations(context), start=1):
        letter = chr(96 + index)
        p = document.add_paragraph(f'({letter}) {declaration}')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_docx_paragraph_font(p)
    p = document.add_paragraph('Commissions, gratuities, or fees:')
    set_docx_paragraph_font(p)
    commission_table = document.add_table(rows=2, cols=4)
    commission_table.style = 'Table Grid'
    for idx, heading in enumerate(['Name of Recipient', 'Address', 'Reason', 'Amount']):
        commission_table.rows[0].cells[idx].text = heading
    for idx, value in enumerate(['None', '-', '-', '-']):
        commission_table.rows[1].cells[idx].text = value
    signature_table = document.add_table(rows=1, cols=2)
    signature_table.style = 'Table Grid'
    signature_table.rows[0].cells[0].text = 'Name'
    signature_table.rows[0].cells[1].text = 'To be completed'
    for field, value in [
        ('In the capacity of', context['signatory_capacity']),
        ('Signed', '____________________________'),
        ('Duly authorized to sign the Bid for and on behalf of', context['company']),
        ('Date', context['date']),
    ]:
        cells = signature_table.add_row().cells
        cells[0].text = field
        cells[1].text = str(value)


def add_pdf_price_schedule_form(elements, bid_document, style):
    bid_pack = bid_document.bid_pack
    references = tender_reference_numbers(bid_pack.tender)
    form_style = zppa_form_styles(style)
    elements.append(Paragraph('Price Schedule: Goods', form_style['title']))
    elements.append(Spacer(1, 8))

    small = style['BodyText'].clone('PriceScheduleSmall')
    small.fontName = 'Times-Roman'
    small.fontSize = 7
    small.leading = 8
    small.alignment = 1
    tiny = small.clone('PriceScheduleTiny')
    tiny.fontSize = 6
    tiny.leading = 7
    tiny.alignment = 1
    bold = small.clone('PriceScheduleBold')
    bold.fontName = 'Times-Bold'

    reference_block = [
        Paragraph(f'Date: {timezone.localdate().strftime("%d/%m/%Y")}', small),
        Paragraph(f'PBN No.: {references["tender_number"]}', small),
        Paragraph('Alternative No.: Not applicable', small),
        Paragraph('Page No.: ____ of ____', small),
    ]
    top_table = Table(
        [['', reference_block]],
        colWidths=[285, 190],
        rowHeights=[72],
        style=price_schedule_outer_style(),
    )
    elements.append(top_table)

    number_row = ['1', '2', '3', '4', '5', '6', '7']
    header_row = [
        Paragraph('Line Item<br/>No', bold),
        Paragraph('Description of Goods', bold),
        Paragraph('Delivery Date', bold),
        Paragraph('Quantity and<br/>physical unit', bold),
        Paragraph('Unit price DDP', bold),
        Paragraph('Total price per line item<br/>(Col. 4 x 5)', bold),
        Paragraph('Country of Origin', bold),
    ]
    instruction_row = [
        Paragraph('[insert<br/>number of<br/>the item]', tiny),
        Paragraph('[insert name of Good]', tiny),
        Paragraph('[insert quoted<br/>Delivery Date]', tiny),
        Paragraph('[insert number of units to<br/>be supplied and name of<br/>the physical unit]', tiny),
        Paragraph('[insert unit price]', tiny),
        Paragraph('[insert total price per line<br/>item]', tiny),
        '',
    ]
    blank_rows = [['', '', '', '', '', '', ''] for _ in range(9)]
    table_rows = [number_row, header_row, instruction_row] + blank_rows
    table = Table(
        table_rows,
        colWidths=[40, 90, 55, 75, 65, 85, 65],
        rowHeights=[12, 34, 34] + [21] * len(blank_rows),
        style=price_schedule_table_style(),
    )
    elements.append(table)
    total_table = Table(
        [[Paragraph('Total Price: Goods', bold), '']],
        colWidths=[185, 100],
        rowHeights=[24],
        style=price_schedule_total_style(),
    )
    total_table.hAlign = 'CENTER'
    elements.append(total_table)


def add_pdf_power_of_attorney_form(elements, bid_document, style):
    bid_pack = bid_document.bid_pack
    tender = bid_pack.tender
    company = bid_pack.company
    references = tender_reference_numbers(tender)
    form_style = zppa_form_styles(style)
    elements.append(Paragraph('POWER OF ATTORNEY / AUTHORISATION OF SIGNATORY', form_style['title']))
    elements.append(Spacer(1, 8))
    reference_rows = plain_field_rows([
        ('Date', timezone.localdate().strftime('%d/%m/%Y')),
        ('Tender / Procurement No.', references['tender_number']),
        ('Tender title', tender.title),
        ('Procuring entity', tender.procuring_entity or '-'),
        ('Bidder', company.name),
    ], style)
    reference_table = Table(reference_rows, colWidths=[150, 220], style=plain_field_table_style())
    reference_table.hAlign = 'RIGHT'
    elements.append(reference_table)
    elements.append(Spacer(1, 10))

    paragraphs = [
        (
            f'We, {company.name}, being a duly registered company participating in the above tender, '
            'hereby appoint the person named below as our true and lawful authorised representative for purposes of this bid.'
        ),
        (
            'The authorised representative is empowered to sign, initial, submit, clarify, and otherwise commit the bidder '
            'in all documents and actions required for this tender, including the Bid Submission Form, Bid-Securing Declaration, '
            'price schedules, declarations, undertakings, and any related correspondence with the Procuring Entity.'
        ),
        (
            'This authority applies only to the tender identified above and remains valid for the bid validity period and any '
            'extension accepted by the bidder, unless revoked in writing by the company.'
        ),
    ]
    for paragraph in paragraphs:
        elements.append(Paragraph(paragraph, form_style['body']))

    elements.append(Spacer(1, 8))
    signatory_rows = plain_field_rows([
        ('Authorised representative name', '________________________________________'),
        ('NRC / Passport No.', '________________________________________'),
        ('Position / capacity', '________________________________________'),
        ('Specimen signature', '________________________________________'),
    ], style)
    elements.append(Table(signatory_rows, colWidths=[190, 285], style=plain_field_table_style()))
    elements.append(Spacer(1, 10))
    company_rows = plain_field_rows([
        ('For and on behalf of', company.name),
        ('Director / Company Secretary name', '________________________________________'),
        ('Position', '________________________________________'),
        ('Signature and company stamp', '________________________________________'),
        ('Witness / Legal Practitioner / Notary', '________________________________________'),
        ('Date', timezone.localdate().strftime('%d/%m/%Y')),
    ], style)
    elements.append(Table(company_rows, colWidths=[190, 285], style=plain_field_table_style()))


def plain_field_rows(rows, style):
    return [
        [
            Paragraph(str(label), style['BodyText']),
            Paragraph(str(value), style['BodyText']),
        ]
        for label, value in rows
    ]


def add_pdf_manufacturer_authorisation_form(elements, bid_pack, style, source_requirement=''):
    tender = bid_pack.tender
    company = bid_pack.company
    references = tender_reference_numbers(tender)
    form_style = zppa_form_styles(style)
    bidder = company.name
    entity = tender.procuring_entity or '[insert complete name of Procuring Entity]'
    bid_no = references['tender_number']
    elements.append(Paragraph("Manufacturer's Authorization", form_style['title']))
    elements.append(Paragraph(
        '[The Bidder shall require the Manufacturer to fill in this Form in accordance with the instructions indicated. '
        'This letter of authorization should be on the letterhead of the Manufacturer and should be signed by a person with '
        'the proper authority to sign documents that are binding on the Manufacturer. The Bidder shall include it in its bid, '
        'if so indicated in the BDS.]',
        form_style['italic'],
    ))
    elements.append(Spacer(1, 10))
    for line in [
        f'Date: {timezone.localdate().strftime("%d/%m/%Y")}',
        f'OIB No.: {bid_no}',
        'Alternative No.: Not applicable',
    ]:
        elements.append(Paragraph(line, form_style['right']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f'To: {entity}', form_style['body']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('WHEREAS', form_style['body']))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(
        f'We, <i>[insert complete name of Manufacturer]</i>, who are official manufacturers of '
        f'<i>[insert type of goods manufactured]</i>, having factories at <i>[insert full address of Manufacturer\'s factories]</i>, '
        f'do hereby authorize {bidder} to submit a bid the purpose of which is to provide the following Goods, '
        f'manufactured by us <i>[insert name and/or brief description of the Goods]</i>, and to subsequently negotiate and sign the Contract.',
        form_style['body'],
    ))
    elements.append(Paragraph(
        'We hereby extend our full guarantee and warranty in accordance with Clause 28 of the General Conditions of Contract, '
        'with respect to the Goods offered by the above firm.',
        form_style['body'],
    ))
    elements.append(Spacer(1, 10))
    for line in [
        'Signed: ________________________________________',
        'Name: __________________________________________',
        'Title: _________________________________________',
    ]:
        elements.append(Paragraph(line, form_style['body']))
        elements.append(Spacer(1, 4))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph('Dated on __________ day of ____________________, ________', form_style['body']))


def bid_securing_declaration_context(bid_pack):
    tender = bid_pack.tender
    company = bid_pack.company
    references = tender_reference_numbers(tender)
    return {
        'date': timezone.localdate().strftime('%d/%m/%Y'),
        'ref_no': references['tender_number'],
        'alternative_no': 'Not applicable',
        'procuring_entity': tender.procuring_entity or '-',
        'bidder': company.name,
        'signatory': '____________________________',
        'capacity': 'Authorised Representative',
        'signing_date': timezone.localdate().strftime('%d/%m/%Y'),
    }


def bid_securing_declaration_paragraphs(context):
    return [
        ('1.', 'We understand that, according to your conditions, bids must be supported by a Bid-Securing Declaration.'),
        ('2.', 'We accept that we shall be liable to suspension from participating in public procurement in accordance with sections 95 and 96 of the Public Procurement Act No. 8 of 2020, if we are in breach of our obligation(s) under the bid conditions because we-'),
        ('(a)', 'have withdrawn our Bid during the period of bid validity specified by us in the Bidding Data Sheet; or'),
        ('(b)', 'having been notified of the acceptance of our Bid by the Procuring Entity during the period of bid validity -'),
        ('(i)', 'fail or refuse to execute the Contract, if required; or'),
        ('(ii)', 'fail or refuse to furnish the Performance Security, in accordance with the Instructions to Bidders.'),
        ('3.', 'We understand this Bid Securing Declaration shall expire if we are not the successful Bidder, on the earlier of'),
        ('(a)', 'our receipt of a copy of your notification of the name of the successful Bidder; or'),
        ('(b)', 'twenty-eight days after the expiration of our Bid.'),
        ('4.', 'We understand that if we are a Joint Venture, the Bid Securing Declaration shall be in the name of the Joint Venture that submits the bid. If the Joint Venture has not been legally constituted at the time of bidding, the Bid Securing Declaration shall be in the names of all future partners as named in the letter of intent.'),
    ]


def add_pdf_bid_securing_declaration(elements, bid_pack, style):
    context = bid_securing_declaration_context(bid_pack)
    form_style = zppa_form_styles(style)
    elements.append(Paragraph('FORM V', form_style['right']))
    elements.append(Paragraph('Reg 88 (2)', form_style['right_small']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('BID-SECURING DECLARATION', form_style['title']))
    elements.append(Spacer(1, 8))
    for line in bid_securing_reference_lines(context):
        elements.append(Paragraph(line, form_style['right']))
    elements.append(Spacer(1, 8))
    for line in bid_securing_recipient_lines(context):
        elements.append(Paragraph(line, form_style['body']))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph('We, the undersigned, declare that:', form_style['body']))
    for marker, text in bid_securing_declaration_paragraphs(context):
        indent = 24 if marker.startswith('(') else 0
        elements.append(Paragraph(f'{marker} {text}', form_style['indented'] if indent else form_style['body']))
    elements.append(Spacer(1, 12))
    for line in bid_securing_signature_lines(context):
        elements.append(Paragraph(line, form_style['body']))


def zppa_form_styles(style):
    body = style['BodyText'].clone('ZppaFormBody')
    body.fontName = 'Times-Roman'
    body.fontSize = 11.5
    body.leading = 15
    body.alignment = 4
    body.spaceAfter = 6

    title = body.clone('ZppaFormTitle')
    title.fontName = 'Times-Bold'
    title.fontSize = 13
    title.leading = 16
    title.alignment = 1
    title.spaceAfter = 8

    right = body.clone('ZppaFormRight')
    right.alignment = 2
    right.spaceAfter = 0

    right_small = right.clone('ZppaFormRightSmall')
    right_small.fontSize = 8
    right_small.leading = 10

    indented = body.clone('ZppaFormIndented')
    indented.leftIndent = 22
    indented.firstLineIndent = -14

    italic = body.clone('ZppaFormItalic')
    italic.fontName = 'Times-Italic'
    italic.fontSize = 9
    italic.leading = 11
    italic.alignment = 4

    return {
        'body': body,
        'title': title,
        'right': right,
        'right_small': right_small,
        'indented': indented,
        'italic': italic,
    }


def bid_securing_field_lines(context):
    return bid_securing_reference_lines(context) + bid_securing_recipient_lines(context)


def bid_securing_reference_lines(context):
    return [
        f'Date: {context["date"]}',
        f'Ref No.: {context["ref_no"]}',
        f'Alternative No.: {context["alternative_no"]}',
    ]


def bid_securing_recipient_lines(context):
    return [
        f'To: {context["procuring_entity"]}',
    ]


def bid_securing_signature_lines(context):
    bidder = escape(context["bidder"])
    return [
        f'Signed: ________________________________ In the capacity of {context["capacity"]}',
        f'Name: {context["signatory"]}',
        f'Duly authorised to sign the bid for and on behalf of: <b>{bidder}</b>',
        f'Dated on {context["signing_date"]}',
    ]


def is_letter_of_bid_requirement(text):
    text = normalize_match_text(text)
    return any(phrase in text for phrase in [
        'bid submission form',
        'letter of bid',
        'form of bid',
        'bid form',
        'service provider bid form',
        'service provider s bid form',
        'signed service provider',
    ])


def is_price_schedule_requirement(text):
    text = normalize_match_text(text)
    return any(phrase in text for phrase in [
        'price schedule',
        'priced schedule',
        'financial offer',
        'bill of quantities',
        'boq',
        'commercial offer',
        'signed quotation',
        'submit a quotation',
    ])


def add_docx_bid_securing_declaration(document, bid_pack):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    context = bid_securing_declaration_context(bid_pack)
    p = document.add_paragraph('FORM V')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_docx_paragraph_font(p, size=12)
    p = document.add_paragraph('Reg 88 (2)')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_docx_paragraph_font(p, size=8)
    p = document.add_paragraph('BID-SECURING DECLARATION')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_paragraph_font(p, size=13, bold=True)
    for line in bid_securing_reference_lines(context):
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_docx_paragraph_font(p)
    for line in bid_securing_recipient_lines(context):
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_docx_paragraph_font(p)
    p = document.add_paragraph('We, the undersigned, declare that:')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_docx_paragraph_font(p)
    for marker, text in bid_securing_declaration_paragraphs(context):
        p = document.add_paragraph(f'{marker} {text}')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if marker.startswith('('):
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-10)
        set_docx_paragraph_font(p)
    for line in bid_securing_signature_lines(context):
        if line.startswith('Duly authorised to sign the bid for and on behalf of:'):
            p = document.add_paragraph()
            prefix = 'Duly authorised to sign the bid for and on behalf of: '
            p.add_run(prefix)
            p.add_run(context['bidder']).bold = True
        else:
            p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_docx_paragraph_font(p)


def set_docx_paragraph_font(paragraph, size=12, bold=False):
    from docx.shared import Pt

    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        if bold:
            run.bold = True


def price_schedule_rows(bid_pack):
    return []


def generate_bid_pack_pdf(bid_pack):
    style = styles()
    settings = SystemSettings.load()
    elements = []
    add_pdf_cover_page(elements, bid_pack)
    elements.append(PageBreak())
    add_pdf_table_of_contents(elements, bid_pack)
    elements.append(PageBreak())

    sections = build_bid_sections(bid_pack)
    for number, (title, lines) in enumerate(sections, start=1):
        if number > 1:
            elements.append(PageBreak())
        add_pdf_part_heading(elements, number, title, bid_pack)
        for line in lines:
            elements.append(Paragraph(str(line), style['BodyText']))
        elements.append(Spacer(1, 6))

        if title == 'Tender and Bidder Information':
            add_pdf_tender_bidder_information(elements, bid_pack, style)

        if title == 'Submission Readiness Gaps':
            rows = [['Item', 'Status', 'Recommended action']]
            gap_rows = submission_gap_rows(bid_pack)
            if gap_rows:
                rows.extend([[item, status, Paragraph(action, style['BodyText'])] for item, status, action in gap_rows])
            else:
                rows.append(['No critical gaps detected', 'Ready', 'Continue final review before submission.'])
            elements.append(Table(rows, colWidths=[165, 80, 230], style=table_style()))
            elements.append(Spacer(1, 6))

        if title == 'Documents Comprising the Bid':
            add_pdf_documents_comprising_bid(elements, bid_pack, style)

        if title == 'XML Envelope Submission Order':
            add_pdf_xml_bid_structure(elements, bid_pack, style, include_heading=False)

        if title == 'Solicitation Forms and Statements':
            add_pdf_required_form_tables(elements, bid_pack, style)

        if title == 'Proof Documents / Eligibility Evidence':
            elements.append(Table([['Document', 'Status'], *company_document_checklist(bid_pack.company)], colWidths=[260, 160], style=table_style()))
            elements.append(Spacer(1, 6))

        if title == 'Qualification Schedules':
            add_pdf_qualification_schedule_summary(elements, bid_pack, style)

        if title == 'Technical Response':
            add_pdf_envelope_response_summary(elements, bid_pack, style, 'technical')

        if title == 'Commercial Offer / Price Schedule':
            add_pdf_envelope_response_summary(elements, bid_pack, style, 'commercial')

    add_signature(elements, label='Authorised signatory')
    main_pdf = build_pdf_response(elements, f'Bid Pack - {bid_pack.tender.title}')
    return append_company_certificate_pdfs(main_pdf, bid_pack)


def generate_bid_pack_docx(bid_pack):
    from docx import Document

    settings = SystemSettings.load()
    currency = settings.default_currency or 'ZMW'
    document = Document()
    apply_formal_docx_styles(document)
    add_docx_cover_page(document, bid_pack)
    document.add_page_break()
    add_docx_table_of_contents(document, bid_pack)
    document.add_page_break()

    for number, (title, lines) in enumerate(build_bid_sections(bid_pack), start=1):
        document.add_heading(f'{number}. {title}', level=1)
        for line in lines:
            document.add_paragraph(docx_safe_text(line), style='List Bullet' if title.endswith('Checklist') else None)

        if title == 'Documents Comprising the Bid':
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for idx, heading in enumerate(['No.', 'Envelope / part', 'Document or requirement', 'Prepared as']):
                table.rows[0].cells[idx].text = heading
            for item in submission_order_rows(bid_pack):
                cells = table.add_row().cells
                cells[0].text = docx_safe_text(item['order'])
                cells[1].text = docx_safe_text(item['envelope'])
                cells[2].text = docx_safe_text(item['requirement'])
                cells[3].text = docx_safe_text(item['prepared_as'])

        if title == 'ITB Ordered Bid Checklist':
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            headings = ['ITB Ref', 'Required Item', 'TenderAI Output', 'Status'] if ordered_itb_rows(bid_pack) else ['Type', 'Requirement', 'Mandatory', 'Action']
            for idx, heading in enumerate(headings):
                table.rows[0].cells[idx].text = heading
            row_source = ordered_itb_rows(bid_pack) or requirements_matrix_rows(bid_pack)
            for first, second, third, fourth in row_source:
                cells = table.add_row().cells
                cells[0].text = docx_safe_text(first)
                cells[1].text = docx_safe_text(second)
                cells[2].text = docx_safe_text(third)
                cells[3].text = docx_safe_text(fourth)

        if title == 'XML Envelope Submission Order':
            add_docx_xml_bid_structure(document, bid_pack, include_heading=False)

        if title == 'Submission Readiness Gaps':
            gap_table = document.add_table(rows=1, cols=3)
            gap_table.style = 'Table Grid'
            for idx, heading in enumerate(['Item', 'Status', 'Recommended action']):
                gap_table.rows[0].cells[idx].text = heading
            rows = submission_gap_rows(bid_pack) or [('No critical gaps detected', 'Ready', 'Continue final review before submission.')]
            for item, status, action in rows:
                cells = gap_table.add_row().cells
                cells[0].text = docx_safe_text(item)
                cells[1].text = docx_safe_text(status)
                cells[2].text = docx_safe_text(action)

        if title == 'Solicitation Forms and Statements':
            add_docx_required_form_tables(document, bid_pack)

        if title == 'Proof Documents / Eligibility Evidence':
            checklist_table = document.add_table(rows=1, cols=2)
            checklist_table.style = 'Table Grid'
            checklist_table.rows[0].cells[0].text = 'Document'
            checklist_table.rows[0].cells[1].text = 'Status'
            for name, status in company_document_checklist(bid_pack.company):
                cells = checklist_table.add_row().cells
                cells[0].text = docx_safe_text(name)
                cells[1].text = docx_safe_text(status)

    document.add_paragraph('\nAuthorised signatory: ________________________')
    document.add_paragraph(f'Prepared by: {settings.default_prepared_by_name or "________________________"}')
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def apply_formal_docx_styles(document):
    from docx.shared import Pt

    styles_collection = document.styles
    for style_name in ['Normal', 'Body Text']:
        try:
            style = styles_collection[style_name]
        except KeyError:
            continue
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
    for style_name, size in [('Title', 18), ('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 11)]:
        try:
            style = styles_collection[style_name]
        except KeyError:
            continue
        style.font.name = 'Times New Roman'
        style.font.size = Pt(size)
        style.font.bold = True


def add_pdf_xml_bid_structure(elements, bid_pack, style, include_heading=True):
    if include_heading:
        elements.append(Paragraph('XML Structured Bid Response', style['Heading1']))
    elements.append(Paragraph(
        'This section follows the Tender Structure XML from ZPPA. Complete and attach the items below in the same envelope order.',
        style['BodyText'],
    ))
    elements.append(Spacer(1, 8))
    for group in xml_bid_structure_groups(bid_pack):
        elements.append(Paragraph(group['title'], style['Heading2']))
        rows = [['No.', 'Requirement', 'Expected response / attachment line', 'TenderAI preparation']]
        for item in group['items']:
            response_items = xml_response_lines(item) or ['To be completed / attached']
            requirement = item.get('requirement') or item.get('title') or '-'
            preparation = prepared_output_for_xml_item(bid_pack, item)
            for line_index, response in enumerate(response_items, start=1):
                rows.append([
                    str(item.get('order') or ''),
                    Paragraph(requirement if line_index == 1 else '', style['BodyText']),
                    Paragraph(response, style['BodyText']),
                    Paragraph(preparation, style['BodyText']),
                ])
        elements.append(Table(rows, colWidths=[34, 150, 205, 86], style=form_table_style()))
        elements.append(Spacer(1, 10))


def add_docx_xml_bid_structure(document, bid_pack, include_heading=True):
    if include_heading:
        document.add_heading('XML Structured Bid Response', level=1)
    document.add_paragraph(
        'This section follows the Tender Structure XML from ZPPA. Complete and attach the items below in the same envelope order.'
    )
    for group in xml_bid_structure_groups(bid_pack):
        document.add_heading(group['title'], level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for idx, heading in enumerate(['No.', 'Requirement', 'Mandatory', 'Expected response / attachment']):
            table.rows[0].cells[idx].text = heading
        for item in group['items']:
            cells = table.add_row().cells
            cells[0].text = docx_safe_text(item.get('order') or '')
            cells[1].text = docx_safe_text(item.get('requirement') or item.get('title') or '-')
            cells[2].text = 'Yes' if item.get('mandatory') else 'No'
            response_items = item.get('response_items') or []
            if response_items:
                cells[3].text = ''
                for response in response_items:
                    cells[3].add_paragraph(docx_safe_text(response), style='List Bullet')
            else:
                cells[3].text = docx_safe_text(item.get('response') or 'To be completed / attached')


def add_pdf_part_heading(elements, number, title, bid_pack=None):
    style = styles()
    elements.append(Paragraph(f'PART {number}', style['SmallMuted']))
    elements.append(Paragraph(title.upper(), style['Heading1']))
    elements.append(Spacer(1, 4))


def add_pdf_tender_bidder_information(elements, bid_pack, style):
    tender = bid_pack.tender
    company = bid_pack.company
    rows = [
        ['Tender title', Paragraph(tender.title, style['BodyText'])],
        ['Procuring entity', tender.procuring_entity or '-'],
        ['Tender number / unique ID', tender.tender_number or '-'],
        ['ZPPA resource ID', tender.zppa_resource_id or '-'],
        ['Procurement method / procedure', tender.procurement_method or '-'],
        ['Submission method', tender.submission_method or '-'],
        ['Bid security', tender.bid_security_amount or 'To confirm from solicitation document'],
        ['Participation fee', tender.participation_fee or '-'],
        ['Closing date', tender.closing_at or tender.closing_date or '-'],
        ['Bidder', company.name],
        ['TPIN', company.tpin or '-'],
        ['PACRA registration', company.registration_number or '-'],
        ['Address', Paragraph(company.address or '-', style['BodyText'])],
        ['Email / phone', f'{company.email or "-"} / {company.phone or "-"}'],
    ]
    elements.append(Table(rows, colWidths=[165, 310], style=form_table_style()))


def add_pdf_documents_comprising_bid(elements, bid_pack, style):
    rows = [['No.', 'Envelope / part', 'Document or requirement', 'Prepared as']]
    ordered_rows = submission_order_rows(bid_pack)
    if ordered_rows:
        rows.extend([
            [
                str(item['order']),
                Paragraph(item['envelope'], style['BodyText']),
                Paragraph(item['requirement'], style['BodyText']),
                Paragraph(item['prepared_as'], style['BodyText']),
            ]
            for item in ordered_rows
        ])
    else:
        rows.extend([
            [str(index), kind, Paragraph(description, style['BodyText']), action]
            for index, (kind, description, mandatory, action) in enumerate(requirements_matrix_rows(bid_pack), start=1)
        ])
    elements.append(Table(rows, colWidths=[34, 120, 210, 111], style=form_table_style()))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
        'Use this table as the assembly checklist. The generated individual bid documents follow the same order.',
        style['FormInstruction'],
    ))


def add_pdf_qualification_schedule_summary(elements, bid_pack, style):
    specs = qualification_form_specs(bid_pack)
    if specs:
        rows = [['No.', 'Qualification form', 'Factor / purpose']]
        for index, spec in enumerate(specs, start=1):
            rows.append([
                str(index),
                Paragraph(spec['title'], style['BodyText']),
                Paragraph(spec.get('factor', 'Qualification'), style['BodyText']),
            ])
        elements.append(Table(rows, colWidths=[35, 250, 190], style=form_table_style()))
        return

    qualification_rows = [
        row for row in submission_order_rows(bid_pack)
        if any(word in normalize_match_text(row['requirement']) for word in ['experience', 'financial', 'turnover', 'personnel', 'equipment', 'eligibility'])
    ]
    if qualification_rows:
        rows = [['No.', 'Qualification requirement', 'Prepared response']]
        rows.extend([
            [str(row['order']), Paragraph(row['requirement'], style['BodyText']), Paragraph(row['prepared_as'], style['BodyText'])]
            for row in qualification_rows
        ])
        elements.append(Table(rows, colWidths=[35, 285, 155], style=form_table_style()))
    else:
        elements.append(Paragraph('No specific qualification schedules were detected yet. Review the solicitation document for eligibility, financial, personnel, equipment, and experience tables.', style['BodyText']))


def add_pdf_envelope_response_summary(elements, bid_pack, style, kind):
    keywords = ['technical', 'specification', 'compliance', 'capacity', 'methodology', 'equipment', 'personnel'] if kind == 'technical' else ['commercial', 'price', 'payment', 'delivery', 'warranty', 'financial', 'offer']
    rows = []
    for item in submission_order_rows(bid_pack):
        text = normalize_match_text(f"{item['envelope']} {item['requirement']} {item['response']} {item['prepared_as']}")
        if any(keyword in text for keyword in keywords):
            rows.append(item)

    if not rows:
        message = 'No technical response items were detected from the XML/ITB yet.' if kind == 'technical' else 'No commercial offer items were detected from the XML/ITB yet.'
        elements.append(Paragraph(message, style['BodyText']))
        return

    table_rows = [['No.', 'Requirement', 'Expected response', 'Prepared as']]
    table_rows.extend([
        [
            str(row['order']),
            Paragraph(row['requirement'], style['BodyText']),
            Paragraph(row['response'], style['BodyText']),
            Paragraph(row['prepared_as'], style['BodyText']),
        ]
        for row in rows
    ])
    elements.append(Table(table_rows, colWidths=[35, 180, 170, 90], style=form_table_style()))


def docx_safe_text(value):
    text = str(value or '')
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)


def add_pdf_cover_page(elements, bid_pack):
    if is_raised_right(bid_pack.company):
        add_raised_right_pdf_cover_page(elements, bid_pack)
        return

    style = styles()
    company = bid_pack.company
    tender = bid_pack.tender
    elements.extend(letterhead_elements(company, 'Bid Submission Pack'))
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(f'<b>Tender:</b> {tender.title}', style['Heading2']))
    elements.append(Paragraph(f'<b>Procuring Entity:</b> {tender.procuring_entity}', style['BodyText']))
    elements.append(Paragraph(f'<b>Tender Number:</b> {tender.tender_number or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Tender Unique ID:</b> {tender.tender_number or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>ZPPA Resource ID:</b> {tender.zppa_resource_id or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Procurement Method:</b> {tender.procurement_method or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Submission Method:</b> {tender.submission_method or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Closing Date:</b> {tender.closing_at or tender.closing_date or "-"}', style['BodyText']))
    elements.append(Spacer(1, 26))
    elements.append(Paragraph(f'<b>Submitted by:</b> {company.name}', style['BodyText']))
    elements.append(Paragraph(f'<b>TPIN:</b> {company.tpin or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Address:</b> {company.address or "-"}', style['BodyText']))
    elements.append(Paragraph(f'<b>Email / Phone:</b> {company.email or "-"} / {company.phone or "-"}', style['BodyText']))


def is_raised_right(company):
    return company.name.strip().lower() == 'raised right investments limited'


def add_raised_right_pdf_cover_page(elements, bid_pack):
    style = styles()
    title_style = style['Title'].clone('RaisedRightCoverTitle')
    title_style.alignment = 1
    title_style.fontSize = 14
    title_style.leading = 18
    title_style.spaceAfter = 0

    label_style = style['BodyText'].clone('RaisedRightCoverLabel')
    label_style.alignment = 1
    label_style.fontName = 'Helvetica-Bold'
    label_style.fontSize = 12
    label_style.leading = 16

    body_style = style['BodyText'].clone('RaisedRightCoverBody')
    body_style.alignment = 1
    body_style.fontSize = 12
    body_style.leading = 18

    detail_style = style['BodyText'].clone('RaisedRightCoverDetail')
    detail_style.alignment = 1
    detail_style.fontSize = 9
    detail_style.leading = 12

    submitted_to = bid_pack.tender.procuring_entity or ''
    tender = bid_pack.tender
    elements.append(Spacer(1, 24))
    elements.append(Paragraph('<b>TENDER</b>', title_style))
    elements.append(Spacer(1, 38))
    elements.append(Paragraph(tender.title, body_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f'<b>Tender Number / Unique ID:</b> {tender.tender_number or "-"}', detail_style))
    elements.append(Paragraph(f'<b>ZPPA Resource ID:</b> {tender.zppa_resource_id or "-"}', detail_style))
    elements.append(Paragraph(f'<b>Procurement Method:</b> {tender.procurement_method or "-"}', detail_style))
    elements.append(Paragraph(f'<b>Submission Method:</b> {tender.submission_method or "-"}', detail_style))
    elements.append(Paragraph(f'<b>Closing Date:</b> {tender.closing_at or tender.closing_date or "-"}', detail_style))
    elements.append(Spacer(1, 54))
    elements.append(Paragraph('<b>SUBMITTED TO:</b>', label_style))
    if submitted_to:
        elements.append(Spacer(1, 14))
        elements.append(Paragraph(submitted_to, body_style))
    elements.append(Spacer(1, 118 if submitted_to else 146))
    elements.append(Paragraph('<b>SUBMITTED BY:</b>', label_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph('Raised Right Investments Ltd', body_style))
    elements.append(Paragraph('Plot No. 31C, Avondale, Off Great East Road', body_style))
    elements.append(Paragraph('Lusaka, Zambia', body_style))
    elements.append(Paragraph('Email: raisedright.zm@gmail.com', body_style))


def add_pdf_table_of_contents(elements, bid_pack):
    style = styles()
    elements.append(Paragraph('Table of Contents', style['Heading1']))
    rows = [['No.', 'Section']]
    rows.extend([[str(index), title] for index, title in enumerate(table_of_contents(bid_pack), start=1)])
    elements.append(Table(rows, colWidths=[45, 390], style=table_style()))


def append_company_certificate_pdfs(main_pdf, bid_pack):
    pdf_documents = ordered_certificate_documents(bid_pack)
    if not pdf_documents:
        return main_pdf

    writer = PdfWriter()
    append_pdf_bytes(writer, main_pdf)
    append_pdf_bytes(writer, attachment_index_pdf(bid_pack, pdf_documents))
    for index, document in enumerate(pdf_documents, start=1):
        try:
            divider = certificate_divider_pdf(document, index)
            append_pdf_bytes(writer, divider)
            with document.file.open('rb') as handle:
                reader = PdfReader(handle)
                for page in reader.pages:
                    writer.add_page(page)
        except Exception:
            continue

    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def ordered_certificate_documents(bid_pack):
    documents = [
        document for document in bid_pack.company.documents.all()
        if document.file and document.file.name.lower().endswith('.pdf')
    ]
    required_types = list(required_document_types(bid_pack.tender))

    def order_key(document):
        try:
            required_index = required_types.index(document.document_type)
        except ValueError:
            required_index = 100
        try:
            base_index = ATTACHMENT_ORDER.index(document.document_type)
        except ValueError:
            base_index = 999
        expiry = document.expiry_date or document.issue_date
        return (required_index, base_index, document.is_expired, expiry or '', document.title.lower())

    return sorted(documents, key=order_key)


def regeneration_status(bid_pack):
    if not bid_pack.generated_pdf or not bid_pack.generated_docx:
        return {
            'needed': True,
            'reason': 'PDF/DOCX files have not been generated yet.',
            'latest_change': None,
        }
    if not bid_pack.generated_at:
        return {
            'needed': True,
            'reason': 'This bid pack was generated before TenderAI started tracking generation dates.',
            'latest_change': None,
        }

    latest_document = bid_pack.company.documents.order_by('-uploaded_at').first()
    latest_task = bid_pack.tender.bid_tasks.order_by('-updated_at').first()
    candidates = []
    if latest_document:
        candidates.append((latest_document.uploaded_at, f'Company document updated: {latest_document.get_document_type_display()}'))
    if latest_task:
        candidates.append((latest_task.updated_at, f'Bid task updated: {latest_task.title}'))
    if bid_pack.tender.updated_at:
        candidates.append((bid_pack.tender.updated_at, 'Tender details were updated.'))

    if not candidates:
        return {'needed': False, 'reason': 'Generated files are current.', 'latest_change': None}

    latest_change, reason = max(candidates, key=lambda item: item[0])
    return {
        'needed': latest_change > bid_pack.generated_at,
        'reason': reason if latest_change > bid_pack.generated_at else 'Generated files are current.',
        'latest_change': latest_change,
    }


def attachment_index_pdf(bid_pack, documents):
    style = styles()
    elements = [
        Paragraph('Attached Company Certificates', style['Heading1']),
        Paragraph(
            'The following company documents are attached after this index in the order TenderAI recommends for bid review.',
            style['BodyText'],
        ),
        Spacer(1, 10),
    ]
    rows = [['No.', 'Document', 'Title', 'Expiry', 'Status']]
    rows.extend([
        [
            str(index),
            document.get_document_type_display(),
            document.title or document.get_document_type_display(),
            document.expiry_date or '-',
            document.status_label,
        ]
        for index, document in enumerate(documents, start=1)
    ])
    elements.append(Table(rows, colWidths=[35, 150, 165, 70, 70], style=table_style()))
    return build_pdf_response(elements, f'Attachment Index - {bid_pack.company.name}')


def append_pdf_bytes(writer, pdf_bytes):
    reader = PdfReader(BytesIO(pdf_bytes))
    for page in reader.pages:
        writer.add_page(page)


def certificate_divider_pdf(company_document, index=None):
    style = styles()
    title_style = style['Title']
    title_style.alignment = 1
    title_style.fontSize = 28
    title_style.leading = 34
    label = f'Attachment {index}' if index else 'Attachment'
    elements = [
        Spacer(1, A4[1] * 0.32),
        Paragraph(f'<b>{label}</b>', style['Heading1']),
        Spacer(1, 18),
        Paragraph(f'<b>{company_document.get_document_type_display()}</b>', title_style),
        Spacer(1, 14),
        Paragraph(f'Title: {company_document.title or company_document.get_document_type_display()}', style['BodyText']),
        Paragraph(f'Expiry: {company_document.expiry_date or "-"}', style['BodyText']),
        Paragraph(f'Status: {company_document.status_label}', style['BodyText']),
    ]
    return build_pdf_response(elements, f'Attachment - {company_document.title}')


def add_docx_cover_page(document, bid_pack):
    company = bid_pack.company
    tender = bid_pack.tender
    document.add_heading('Bid Submission Pack', 0)
    document.add_paragraph(f'Tender: {tender.title}')
    document.add_paragraph(f'Procuring Entity: {tender.procuring_entity}')
    document.add_paragraph(f'Tender Number / Unique ID: {tender.tender_number or "-"}')
    document.add_paragraph(f'ZPPA Resource ID: {tender.zppa_resource_id or "-"}')
    document.add_paragraph(f'Procurement Method: {tender.procurement_method or "-"}')
    document.add_paragraph(f'Submission Method: {tender.submission_method or "-"}')
    document.add_paragraph(f'Closing Date: {tender.closing_at or tender.closing_date or "-"}')
    document.add_paragraph('')
    document.add_heading('Submitted by', level=1)
    document.add_paragraph(company.name)
    document.add_paragraph(f'TPIN: {company.tpin or "-"}')
    document.add_paragraph(f'Address: {company.address or "-"}')
    document.add_paragraph(f'Email: {company.email or "-"}')
    document.add_paragraph(f'Phone: {company.phone or "-"}')


def add_docx_table_of_contents(document, bid_pack):
    document.add_heading('Table of Contents', level=1)
    for index, title in enumerate(table_of_contents(bid_pack), start=1):
        document.add_paragraph(f'{index}. {title}')


def build_bid_sections(bid_pack):
    tender = bid_pack.tender
    company = bid_pack.company
    sections = [
        ('Tender and Bidder Information', [
            'This front section identifies the tender, procuring entity, bidder, and submission reference information used in the forms that follow.',
        ]),
        ('Submission Readiness Gaps', [
            'Fix the following missing or expired items before treating this bid pack as ready for submission.',
        ]),
        ('Documents Comprising the Bid', [
            'This section converts ITB 11.1 / Tender Structure XML into a practical submission order.',
        ]),
    ]
    if has_xml_bid_structure(bid_pack):
        sections.append((
            'XML Envelope Submission Order',
            ['The response structure below follows the Tender Structure XML envelope order from ZPPA.'],
        ))
    else:
        sections.append((
            'ITB Ordered Bid Checklist',
            ['The checklist below follows the order captured from the solicitation ITB Documents Comprising the Bid table.'],
        ))
    sections.extend([
        ('Solicitation Forms and Statements', [
            'The forms and declarations below are arranged first because most ZPPA bid documents require signed forms before supporting evidence.',
        ]),
        ('Proof Documents / Eligibility Evidence', [
            'Attach mandatory registration, tax, compliance, and other proof documents in the order required by the tender.',
        ]),
        ('Qualification Schedules', [
            'Complete qualification tables such as eligibility, financial situation, personnel, equipment, and experience where required.',
        ]),
        ('Technical Response', [
            'Respond to the technical specifications, compliance requirements, delivery capacity, personnel, equipment, and methodology items.',
        ]),
        ('Commercial Offer / Price Schedule', [
            'Attach or complete the tender price schedule, bill of quantities, payment terms, delivery period, warranty, or financial offer required by the solicitation document.',
        ]),
        ('Company Profile Summary', [
            company.profile_summary or f'{company.name} is a registered supplier/contractor.',
            f'TPIN: {company.tpin or "-"}',
            f'PACRA Registration: {company.registration_number or "-"}',
        ]),
    ])
    return sections


def save_generated_files(bid_pack):
    pdf_bytes = generate_bid_pack_pdf(bid_pack)
    docx_bytes = generate_bid_pack_docx(bid_pack)
    safe_id = bid_pack.pk or 'new'
    bid_pack.generated_pdf.save(f'bid-pack-{safe_id}.pdf', ContentFile(pdf_bytes), save=False)
    bid_pack.generated_docx.save(f'bid-pack-{safe_id}.docx', ContentFile(docx_bytes), save=False)
    bid_pack.generated_at = timezone.now()
    bid_pack.save(update_fields=['generated_pdf', 'generated_docx', 'generated_at'])


def generate_xml_bid_documents(bid_pack):
    created = []
    active_ids = []
    for group in xml_bid_structure_groups(bid_pack):
        for item in group['items']:
            order = int(item.get('order') or len(created) + 1)
            requirement = item.get('requirement') or item.get('title') or 'Bid requirement'
            bid_document, _ = BidDocument.objects.update_or_create(
                bid_pack=bid_pack,
                order=order,
                requirement=requirement,
                defaults={
                    'envelope': group['title'],
                    'expected_response': '\n'.join(item.get('response_items') or [item.get('response', '')]).strip(),
                    'mandatory': bool(item.get('mandatory')),
                    'matched_company_document': match_company_document_for_xml_item(bid_pack, item),
                },
            )
            pdf_bytes = generate_single_bid_document_pdf(bid_document)
            bid_document.generated_pdf.save(
                f'{safe_slug(bid_document.order)}-{safe_slug(bid_document.requirement)}.pdf',
                ContentFile(pdf_bytes),
                save=False,
            )
            bid_document.generated_at = timezone.now()
            bid_document.save(update_fields=['generated_pdf', 'generated_at'])
            created.append(bid_document)
            active_ids.append(bid_document.pk)
    bid_pack.bid_documents.exclude(pk__in=active_ids).delete()
    return created


def generate_single_bid_document_pdf(bid_document):
    style = styles()
    elements = []
    bid_pack = bid_document.bid_pack
    lower = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    form_code = qualification_form_code_for_text(f'{bid_document.requirement} {bid_document.expected_response}')
    extracted_template = extracted_template_for_bid_document(bid_document, form_code)

    if should_generate_simple_instruction_only(bid_document, lower, extracted_template):
        add_pdf_simple_instruction_only_document(elements, bid_document, style)
        return build_pdf_response(elements, bid_document.requirement)

    add_pdf_bid_document_instruction_page(elements, bid_document, style)
    required_type = expected_document_type_for_text(lower)
    if bid_document.matched_company_document and not is_experience_requirement(lower) and (
        looks_like_attachment_request(lower) or required_type
    ):
        pdf = build_pdf_response(elements, bid_document.requirement)
        return append_matched_certificate_pdf(pdf, bid_document)

    elements.append(PageBreak())

    if 'bid declaration' in lower or 'bid securing' in lower or 'bid security' in lower:
        elements.extend(letterhead_elements(bid_pack.company, '', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_bid_securing_declaration(elements, bid_pack, style)
    elif is_letter_of_bid_requirement(lower):
        elements.extend(letterhead_elements(bid_pack.company, '', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_letter_of_bid(elements, bid_pack, style)
        if is_price_schedule_requirement(lower):
            elements.append(PageBreak())
            elements.extend(letterhead_elements(bid_pack.company, '', use_pdf_letterhead=True))
            elements.append(Spacer(1, 8))
            add_pdf_price_schedule_form(elements, bid_document, style)
    elif is_power_of_attorney_requirement(lower):
        elements.extend(letterhead_elements(bid_pack.company, '', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        if can_use_extracted_template_for_bid_document(bid_document, extracted_template):
            add_pdf_extracted_form_template(elements, bid_document, extracted_template, style)
        else:
            add_pdf_power_of_attorney_form(elements, bid_document, style)
    elif form_code:
        if form_code != 'EXP-2.4.1':
            letterhead_title = '' if form_code == 'MFR' else QUALIFICATION_FORM_DEFINITIONS[form_code]['title']
            elements.extend(letterhead_elements(bid_pack.company, letterhead_title, use_pdf_letterhead=True))
            elements.append(Spacer(1, 8))
        add_pdf_qualification_form(elements, bid_pack, form_code, style, source_requirement=bid_document.requirement)
    elif is_experience_requirement(lower):
        elements.extend(letterhead_elements(bid_pack.company, 'Experience Evidence Schedule', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_experience_evidence_form(elements, bid_document, style)
    elif required_type or looks_like_attachment_request(lower):
        elements.pop()
        return build_pdf_response(elements, bid_document.requirement)
    elif 'warranty' in lower or 'delivery period' in lower or 'delivery' in lower:
        elements.extend(letterhead_elements(bid_pack.company, 'Delivery / Warranty Undertaking', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_delivery_warranty_form(elements, bid_document, style)
    elif 'technical' in lower or 'specification' in lower or 'qualification and experience' in lower or 'general experience' in lower:
        elements.extend(letterhead_elements(bid_pack.company, 'Technical Response Schedule', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_technical_response_form(elements, bid_document, style)
    elif bid_document.matched_company_document:
        pdf = build_pdf_response(elements, bid_document.requirement)
        return append_matched_certificate_pdf(pdf, bid_document)
    elif required_type:
        add_pdf_missing_certificate_notice(elements, bid_document, style)
    elif looks_like_table_requirement(lower):
        elements.extend(letterhead_elements(bid_pack.company, 'Required Schedule / Table', use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_required_table_response(elements, bid_document, style)
    elif can_use_extracted_template_for_bid_document(bid_document, extracted_template):
        elements.extend(letterhead_elements(bid_pack.company, form_document_title(bid_document), use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_extracted_form_template(elements, bid_document, extracted_template, style)
    else:
        elements.extend(letterhead_elements(bid_pack.company, form_document_title(bid_document), use_pdf_letterhead=True))
        elements.append(Spacer(1, 8))
        add_pdf_placeholder_form(elements, bid_document, 'Tender Requirement Response', style)

    if not bid_document.matched_company_document and not generated_document_has_own_signature(bid_document):
        add_signature(elements, label='Authorised signatory')
    pdf = build_pdf_response(elements, bid_document.requirement)
    pdf = apply_company_letterhead_pdf(pdf, bid_pack.company, skip_first_page=True)
    return append_matched_certificate_pdf(pdf, bid_document)


def generated_document_has_own_signature(bid_document):
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    return any(phrase in text for phrase in [
        'bid declaration',
        'bid securing',
        'bid security',
        'bid submission form',
        'letter of bid',
        'bid form',
        'form of bid',
        'signed service provider',
        'power of attorney',
        'past experience',
        'similar experience',
        'completion certificate',
        'general experience',
        'specific experience',
        'key activit',
        'manufacturer',
        'manufacturer authorisation',
        'manufacturer authorization',
    ])


def should_generate_simple_instruction_only(bid_document, lower, extracted_template=None):
    if 'power of attorney' in lower:
        return False
    return any(phrase in lower for phrase in [
        'site visit certificate',
        'site visit certificates',
        'site visit',
        'litigation status',
        'legal practitioner',
    ])


def is_power_of_attorney_template(template):
    if not template:
        return False
    if template.get('code') == 'POWER_OF_ATTORNEY':
        return True
    title = normalize_match_text(f"{template.get('title', '')} {template.get('heading', '')}")
    return 'power of attorney' in title and 'form' in title


def is_power_of_attorney_requirement(text):
    return 'power of attorney' in normalize_match_text(text)


def can_use_extracted_template_for_bid_document(bid_document, template):
    if not template:
        return False
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    return template.get('code') == 'POWER_OF_ATTORNEY' and 'power of attorney' in text


def add_pdf_simple_instruction_only_document(elements, bid_document, style):
    bid_pack = bid_document.bid_pack
    title, actions = simple_instruction_details(bid_document)
    elements.extend(letterhead_elements(bid_pack.company, '', use_pdf_letterhead=True))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(title, style['FormTitle']))
    elements.append(Spacer(1, 8))
    rows = [
        ['Tender', Paragraph(bid_pack.tender.title, style['BodyText'])],
        ['Procuring entity', bid_pack.tender.procuring_entity or '-'],
        ['Bidder', bid_pack.company.name],
        ['XML part', bid_document.envelope or '-'],
        ['Requirement', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Instruction', Paragraph('<br/>'.join(bid_document.expected_response_lines) or bid_document.requirement, style['BodyText'])],
    ]
    elements.append(Table(rows, colWidths=[130, 345], style=form_table_style()))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph('What to do', style['Heading3']))
    for action in actions:
        elements.append(Paragraph(f'- {action}', style['BodyText']))


def simple_instruction_details(bid_document):
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    if 'power of attorney' in text:
        return 'POWER OF ATTORNEY / AUTHORISATION INSTRUCTION', [
            'Attach a signed Power of Attorney, board resolution, or authorisation letter for the person signing the bid.',
            'Make sure the authorised signatory name matches the signed bid forms.',
            'Use the official template from the solicitation if one is provided.',
        ]
    if 'litigation' in text or 'legal practitioner' in text:
        return 'LITIGATION STATUS LETTER INSTRUCTION', [
            'Obtain the Litigation Status letter from a qualified legal practitioner where required.',
            'Do not replace this with a bidder-written statement if the tender says it will not be accepted.',
            'Attach the signed and dated letter to this part of the bid.',
        ]
    return 'SITE VISIT CERTIFICATE INSTRUCTION', [
        'Attach the fully signed/stamped site visit certificate issued for this tender.',
        'Confirm the certificate shows the correct tender, site, company, and representative.',
        'Scan and upload the certificate under company/tender documents before final submission.',
    ]


def add_pdf_bid_document_instruction_page(elements, bid_document, style):
    bid_pack = bid_document.bid_pack
    full_text = f'{bid_document.requirement} {bid_document.expected_response}'
    required_type = expected_document_type_for_text(full_text)
    experience_plan = experience_document_plan(full_text)
    matched = bid_document.matched_company_document
    action = bid_document_action_label(bid_document)
    status = 'Ready - company document attached' if matched else ('Missing required certificate' if required_type else ('Experience form to complete' if experience_plan else 'Manual response / form to complete'))
    title = 'CERTIFICATE ATTACHMENT INSTRUCTION' if matched and required_type else ('EXPERIENCE FORM INSTRUCTION' if experience_plan else 'XML BID DOCUMENT INSTRUCTION')
    elements.append(Paragraph(title, style['FormTitle']))
    elements.append(Paragraph(
        bid_document_instruction_intro(bid_document, required_type, matched),
        style['FormInstruction'],
    ))
    elements.append(Spacer(1, 10))
    rows = [
        ['XML item', str(bid_document.order)],
        ['Envelope / part', bid_document.envelope or '-'],
        ['Tender', Paragraph(bid_pack.tender.title, style['BodyText'])],
        ['Bidder', bid_pack.company.name],
        ['Requirement', Paragraph(bid_document.requirement, style['BodyText'])],
        ['TenderAI action', Paragraph(action, style['BodyText'])],
        ['Current status', status],
    ]
    if not is_power_of_attorney_requirement(full_text):
        rows.insert(5, ['Instruction from XML', Paragraph('<br/>'.join(bid_document.expected_response_lines) or 'Complete / attach as required.', style['BodyText'])])
    if required_type:
        rows.append(['Required certificate type', required_type.label])
    if experience_plan:
        rows.append(['Experience form required', experience_plan['form_title']])
        rows.append(['Supporting evidence needed', Paragraph('<br/>'.join(experience_plan['evidence']), style['BodyText'])])
    if matched:
        rows.append(['Matched document', f'{matched.get_document_type_display()} - {matched.title}'])
    attachment_warning = bid_document_attachment_warning(bid_document)
    if attachment_warning:
        rows.append(['Attachment warning', Paragraph(attachment_warning, style['BodyText'])])
    elements.append(Table(rows, colWidths=[145, 330], style=form_table_style()))


def certificate_instruction_text(bid_document):
    document = bid_document.matched_company_document
    return (
        f'TenderAI matched this XML requirement to the uploaded {document.get_document_type_display()} document. '
        'Confirm that the certificate is readable, current, and belongs to the bidding company. The actual uploaded certificate follows immediately after this instruction page.'
    )


def bid_document_instruction_intro(bid_document, required_type, matched):
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    if matched and required_type:
        return certificate_instruction_text(bid_document)
    if is_experience_requirement(text):
        plan = experience_document_plan(text)
        return (
            f'TenderAI detected an experience requirement. Prepare {plan["form_title"]} and attach the supporting evidence listed below. '
            'Do not attach every experience form unless the solicitation specifically requests each one.'
        )
    if required_type or looks_like_attachment_request(text):
        return 'This page explains the required attachment. Upload or attach the correct supporting document before final bid submission.'
    return 'This page explains what must be prepared for this single XML requirement. The next page contains the generated form, letter, table, or response schedule.'


def bid_document_action_label(bid_document):
    if bid_document.matched_company_document:
        return 'Attach the matched company certificate/document after this instruction page.'
    plan = experience_document_plan(f'{bid_document.requirement} {bid_document.expected_response}')
    if plan:
        return f'Complete {plan["form_title"]} and attach: {", ".join(plan["evidence"])}.'
    return prepared_output_for_bid_document(bid_document)


def bid_document_rows_for_display(bid_pack):
    rows = []
    for document in bid_pack.bid_documents.order_by('order'):
        required_type = expected_document_type_for_text(f'{document.requirement} {document.expected_response}')
        attachment_warning = bid_document_attachment_warning(document)
        if document.matched_company_document:
            if document.matched_company_document.is_expired:
                status = 'Expired'
                badge = 'danger'
                action = attachment_warning
            elif attachment_warning:
                status = 'Upgrade required'
                badge = 'warning'
                action = attachment_warning
            else:
                status = 'Ready'
                badge = 'success'
                action = f'{document.matched_company_document.get_document_type_display()} attached'
        elif required_type:
            status = 'Missing'
            badge = 'warning'
            action = f'Upload {required_type.label}'
        else:
            status = 'Prepare'
            badge = 'info'
            action = prepared_output_for_bid_document(document)
        rows.append({
            'document': document,
            'required_type': required_type,
            'status': status,
            'badge': badge,
            'action': action,
            'warning': attachment_warning,
        })
    return rows


def add_pdf_certificate_attachment_notice(elements, bid_document, style):
    document = bid_document.matched_company_document
    elements.append(Spacer(1, A4[1] * 0.22))
    elements.append(Paragraph('CERTIFICATE / DOCUMENT ATTACHED', style['FormTitle']))
    elements.append(Spacer(1, 10))
    rows = [
        ['Required item', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Attached document', document.get_document_type_display()],
        ['Document title', document.title or document.get_document_type_display()],
        ['Expiry date', document.expiry_date or '-'],
        ['Status', document.status_label],
    ]
    elements.append(Table(rows, colWidths=[155, 320], style=form_table_style()))
    elements.append(Paragraph('The actual uploaded certificate/document follows this page.', style['FormInstruction']))


def add_pdf_missing_certificate_notice(elements, bid_document, style):
    required_type = expected_document_type_for_text(f'{bid_document.requirement} {bid_document.expected_response}')
    elements.append(Spacer(1, A4[1] * 0.2))
    elements.append(Paragraph('MISSING REQUIRED CERTIFICATE', style['FormTitle']))
    elements.append(Spacer(1, 10))
    rows = [
        ['Required item', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Expected document type', required_type.label if required_type else 'Certificate / supporting document'],
        ['Instruction', Paragraph('<br/>'.join(bid_document.expected_response_lines) or bid_document.requirement, style['BodyText'])],
        ['Action required', 'Upload this certificate/document under company documents, then regenerate XML bid documents.'],
    ]
    elements.append(Table(rows, colWidths=[155, 320], style=form_table_style()))


def looks_like_table_requirement(text):
    text = normalize_match_text(text)
    return any(word in text for word in ['summary', 'schedule', 'table', 'assignments', 'personnel', 'reference letters', 'financial statements'])


def is_experience_requirement(text):
    text = normalize_match_text(text)
    return any(word in text for word in [
        'experience',
        'past experience',
        'similar experience',
        'completion certificate',
        'completion certificates',
        'reference letter',
        'reference letters',
        'works of similar nature',
        'assignments executed',
        'past contract',
    ])


def experience_required_count(text):
    text = normalize_match_text(text)
    digit_match = re.search(r'\b(\d+)\s+(?:completion certificates|certificates|contracts|assignments|references)\b', text)
    if digit_match:
        return max(1, min(int(digit_match.group(1)), 10))
    word_numbers = {
        'one': 1,
        'two': 2,
        'three': 3,
        'four': 4,
        'five': 5,
    }
    for word, number in word_numbers.items():
        if re.search(rf'\b{word}\b.*\b(?:completion certificates|certificates|contracts|assignments|references)\b', text):
            return number
    return 3


def add_pdf_experience_evidence_form(elements, bid_document, style):
    requirement_text = f'{bid_document.requirement} {bid_document.expected_response}'
    form_code = experience_form_code_for_text(requirement_text)
    required_count = experience_required_count(requirement_text)
    plan = experience_document_plan(requirement_text)
    if form_code == 'EXP-2.4.1':
        add_pdf_general_experience_form(elements, bid_document.bid_pack, style, required_count=required_count)
        add_pdf_experience_supporting_evidence(elements, plan, style)
        elements.append(Spacer(1, 10))
        elements.append(Paragraph('Required by XML:', style['Heading3']))
        for line in bid_document.expected_response_lines or [bid_document.requirement]:
            elements.append(Paragraph(line, style['BodyText']))
        return
    elements.append(Paragraph(experience_form_title(form_code), style['FormTitle']))
    elements.append(Paragraph(
        experience_form_instruction(form_code),
        style['FormInstruction'],
    ))
    rows = [experience_form_headers(form_code)]
    for index in range(1, required_count + 1):
        rows.append(experience_form_blank_row(form_code, index))
    col_width = 475 / len(rows[0])
    elements.append(Table(rows, colWidths=[col_width] * len(rows[0]), style=form_table_style()))
    add_pdf_experience_supporting_evidence(elements, plan, style)
    elements.append(Spacer(1, 10))
    elements.append(Paragraph('Required by XML:', style['Heading3']))
    for line in bid_document.expected_response_lines or [bid_document.requirement]:
        elements.append(Paragraph(line, style['BodyText']))


def experience_form_code_for_text(text):
    text = normalize_match_text(text)
    if 'key activit' in text or 'critical activit' in text:
        return 'EXP-2.4.2(b)'
    if any(phrase in text for phrase in [
        'specific experience',
        'similar nature',
        'similar assignment',
        'similar assignments',
        'similar work',
        'similar works',
        'similar contract',
        'similar contracts',
        'reference letter',
        'reference letters',
        'completion certificate',
        'completion certificates',
    ]):
        return 'EXP-2.4.2(a)'
    if 'general experience' in text or 'past experience' in text:
        return 'EXP-2.4.1'
    return 'EXP-2.4.2(a)'


def experience_document_plan(text):
    if not is_experience_requirement(text):
        return None
    code = experience_form_code_for_text(text)
    evidence = experience_supporting_evidence(text, code)
    return {
        'form_code': code,
        'form_title': QUALIFICATION_FORM_DEFINITIONS.get(code, {}).get('title', experience_form_title(code)),
        'evidence': evidence,
    }


def experience_supporting_evidence(text, code):
    text = normalize_match_text(text)
    evidence = []
    if code == 'EXP-2.4.1':
        evidence.extend([
            'contracts, award letters, LPOs, GRNs, or delivery notes',
            'completion certificates or client confirmations where available',
        ])
    if 'reference letter' in text or 'references' in text:
        evidence.append('traceable client reference letters with contact details')
    if 'completion certificate' in text or 'completion certificates' in text:
        evidence.append('completion certificates')
    if 'contract' in text or 'assignment' in text or 'past experience' in text:
        evidence.append('contracts, award letters, LPOs, GRNs, or delivery notes')
    if 'key activit' in text or code == 'EXP-2.4.2(b)':
        evidence.append('evidence showing the required key activity, quantity, scope, and completion')
    if not evidence:
        evidence.append('similar assignment evidence such as contracts, completion certificates, award letters, GRNs, delivery notes, or reference letters')
    return list(dict.fromkeys(evidence))


def experience_form_title(code):
    return QUALIFICATION_FORM_DEFINITIONS.get(code, {}).get('title', 'Experience Evidence Schedule').upper()


def add_pdf_experience_supporting_evidence(elements, plan, style):
    if not plan:
        return
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('Supporting evidence to attach behind this form:', style['Heading3']))
    for item in plan['evidence']:
        elements.append(Paragraph(f'- {item}', style['BodyText']))


def experience_form_instruction(code):
    if code == 'EXP-2.4.1':
        return 'List general experience and attach evidence such as completion certificates, contracts, award letters, GRNs, or reference letters.'
    if code == 'EXP-2.4.2(b)':
        return 'List specific experience in the required key activities and attach evidence showing quantity, scope, completion, and client confirmation.'
    return 'List similar contracts/assignments and attach completion certificates, award letters, contracts, GRNs, or reference evidence.'


def experience_form_headers(code):
    if code == 'EXP-2.4.1':
        return ['No.', 'Start', 'End', 'Client', 'Contract / assignment', 'Role', 'Evidence']
    if code == 'EXP-2.4.2(b)':
        return ['No.', 'Key activity', 'Contract reference', 'Quantity / scope', 'Completion date', 'Evidence']
    return ['No.', 'Contract / assignment', 'Client', 'Country', 'Completion date', 'Similarity', 'Evidence']


def experience_form_blank_row(code, index):
    if code == 'EXP-2.4.1':
        return [str(index), '____', '____', '________________', '____________________', 'Contractor', 'Attach proof']
    if code == 'EXP-2.4.2(b)':
        return [str(index), '________________', '________________', '________________', '____', 'Attach proof']
    return [str(index), '____________________', '________________', 'Zambia', '____', 'Similar works/services', 'Attach proof']


def add_pdf_general_experience_form(elements, bid_pack, style, required_count=6):
    normal = style['BodyText'].clone('GeneralExperienceNormal')
    normal.fontName = 'Times-Roman'
    normal.fontSize = 11.5
    normal.leading = 15

    small = normal.clone('GeneralExperienceSmall')
    small.fontSize = 9.5
    small.leading = 11

    centered = normal.clone('GeneralExperienceCentered')
    centered.alignment = 1

    title = style['FormTitle'].clone('GeneralExperienceTitle')
    title.fontName = 'Times-Bold'
    title.fontSize = 13
    title.leading = 16
    title.spaceAfter = 2

    elements.append(Paragraph('<b>Experience</b>', title))
    elements.append(Paragraph('<b>General Experience</b>', title))
    elements.append(Spacer(1, 18))

    company = bid_pack.company
    tender = bid_pack.tender
    references = tender_reference_numbers(tender)
    top_rows = [
        [
            Paragraph("Bidder's Legal Name:  ______________________________", normal),
            Paragraph(f'Date:  {timezone.localdate().strftime("%d/%m/%Y")}', normal),
        ],
        [
            Paragraph('JV Partner Legal Name:  ___________________________', normal),
            Paragraph(f'Bidding No.:  {references["tender_number"] or "________________"}', normal),
        ],
        ['', Paragraph('Page  ______ of ______ pages', normal)],
    ]
    if company.name:
        top_rows[0][0] = Paragraph(f"Bidder's Legal Name:  {company.name}", normal)
    elements.append(Table(top_rows, colWidths=[290, 185], style=TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 11.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ])))
    elements.append(Spacer(1, 18))

    header = [
        Paragraph('Starting<br/>Month /<br/>Year', centered),
        Paragraph('Ending<br/>Month /<br/>Year', centered),
        Paragraph('Years*', centered),
        Paragraph('Contract Identification', centered),
        Paragraph('Role of<br/>Bidder', centered),
    ]
    rows = [header]
    count = max(6, min(required_count or 6, 8))
    contract_cell = (
        'Contract name:<br/>'
        'Brief Description of the Works performed by the Bidder:<br/>'
        'Name of Employer:<br/>'
        'Address:'
    )
    for _ in range(count):
        rows.append([
            Paragraph('______', centered),
            Paragraph('______', centered),
            Paragraph('', centered),
            Paragraph(contract_cell, small),
            Paragraph('______', centered),
        ])

    table = Table(rows, colWidths=[55, 58, 42, 255, 65], rowHeights=[52] + [64] * count)
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.55, colors.black),
        ('FONTNAME', (0, 0), (-1, 0), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, 0), 11.5),
        ('LEADING', (0, 0), (-1, 0), 13),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (2, -1), 'CENTER'),
        ('ALIGN', (4, 1), (4, -1), 'CENTER'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(
        '*List calendar year for years with contracts with at least nine (9) months activity per year starting with the earliest year',
        small,
    ))


def add_pdf_required_table_response(elements, bid_document, style):
    rows = [['No.', 'Required information', 'Bidder response / attachment']]
    response_lines = bid_document.expected_response_lines or [bid_document.requirement]
    for index, line in enumerate(response_lines, start=1):
        rows.append([
            str(index),
            Paragraph(line, style['BodyText']),
            Paragraph(default_response_for_requirement(bid_document), style['BodyText']),
        ])
    elements.append(Paragraph('REQUIRED SCHEDULE / TABLE RESPONSE', style['FormTitle']))
    elements.append(Table(rows, colWidths=[35, 240, 200], style=form_table_style()))


def form_document_title(bid_document):
    code = qualification_form_code_for_text(f'{bid_document.requirement} {bid_document.expected_response}')
    if code:
        return QUALIFICATION_FORM_DEFINITIONS[code]['title']
    lower = bid_document.requirement.lower()
    if 'bid declaration' in lower or 'bid securing' in lower or 'bid security' in lower:
        return 'Bid-Securing Declaration'
    if is_letter_of_bid_requirement(lower):
        return 'Letter of Bid'
    return bid_document.requirement


def extracted_template_for_bid_document(bid_document, form_code=None):
    templates = []
    for document in bid_document.bid_pack.tender.solicitation_documents.order_by('-uploaded_at'):
        analysis = document.analysis_summary or {}
        templates.extend(analysis.get('extracted_form_templates', []) or [])
    if not templates:
        return None
    wanted_codes = []
    lower = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    if form_code:
        wanted_codes.append(form_code)
    if 'bid securing' in lower or 'bid security' in lower or 'bid declaration' in lower:
        wanted_codes.append('BID_SECURING_DECLARATION')
    if is_letter_of_bid_requirement(lower):
        wanted_codes.append('LETTER_OF_BID')
    if is_power_of_attorney_requirement(lower):
        wanted_codes.append('POWER_OF_ATTORNEY')
    for template in templates:
        if template.get('code') in wanted_codes:
            return template
    for template in templates:
        title = normalize_match_text(f"{template.get('title', '')} {template.get('heading', '')}")
        if any(word in lower for word in title.split() if len(word) > 5):
            return template
    return None


def add_pdf_extracted_form_template(elements, bid_document, template, style):
    bid_pack = bid_document.bid_pack
    elements.append(Paragraph(str(template.get('heading') or template.get('title') or bid_document.requirement).upper(), style['FormTitle']))
    elements.append(Paragraph(
        'TenderAI used the form structure detected from the uploaded solicitation document. Confirm layout and complete blank fields before submission.',
        style['FormInstruction'],
    ))
    elements.append(Spacer(1, 6))

    fields = template.get('fields') or []
    if fields:
        rows = [['Field from solicitation form', 'TenderAI filled value / bidder response']]
        for field in fields:
            rows.append([
                Paragraph(field, style['BodyText']),
                Paragraph(fill_extracted_form_field(field, bid_pack), style['BodyText']),
            ])
        elements.append(Table(rows, colWidths=[190, 285], style=form_table_style()))
        elements.append(Spacer(1, 8))

    rows = template.get('rows') or []
    if rows:
        normalized_rows = normalize_extracted_rows(rows, style, bid_pack)
        col_count = max(len(row) for row in normalized_rows)
        col_width = 475 / col_count
        elements.append(Paragraph('Original table structure detected:', style['Heading3']))
        elements.append(Table(normalized_rows, colWidths=[col_width] * col_count, style=form_table_style()))
        elements.append(Spacer(1, 8))

    lines = template.get('lines') or []
    if lines:
        elements.append(Paragraph('Original form text / instructions detected:', style['Heading3']))
        for line in lines[:22]:
            if not line or line == template.get('heading'):
                continue
            filled = fill_known_placeholders(line, bid_pack)
            elements.append(Paragraph(filled, style['BodyText']))

    if not fields and not rows and not lines:
        add_pdf_placeholder_form(elements, bid_document, template.get('title') or 'Tender Form Response', style)
        return

    elements.append(Spacer(1, 8))
    elements.append(Table([
        ['Bidder', bid_pack.company.name],
        ['Authorised representative', 'To be completed'],
        ['Signature and stamp', '____________________________'],
        ['Date', timezone.localdate().strftime('%d/%m/%Y')],
    ], colWidths=[180, 295], style=form_table_style()))


def normalize_extracted_rows(rows, style, bid_pack):
    width = max(len(row) for row in rows)
    normalized = []
    for row in rows[:16]:
        normalized.append([
            Paragraph(fill_known_placeholders(str(cell), bid_pack), style['BodyText'])
            for cell in [*row, *([''] * (width - len(row)))]
        ])
    return normalized


def fill_extracted_form_field(field, bid_pack):
    label = normalize_match_text(field)
    company = bid_pack.company
    tender = bid_pack.tender
    if 'bidder' in label and 'name' in label:
        return company.name
    if label in {'name'} or 'legal name' in label:
        return company.name
    if 'address' in label:
        return company.address or 'To be completed'
    if 'country' in label:
        return 'Zambia'
    if 'tpin' in label or 'taxpayer' in label:
        return company.tpin or 'To be completed'
    if 'registration' in label or 'incorporation' in label:
        return company.registration_number or 'To be completed from PACRA certificate'
    if 'email' in label or 'mail' in label:
        return company.email or 'To be completed'
    if 'telephone' in label or 'phone' in label:
        return company.phone or 'To be completed'
    if 'tender' in label or 'contract' in label:
        return tender.title
    if 'procuring entity' in label or label == 'to':
        return tender.procuring_entity or 'To be completed'
    if 'date' in label:
        return timezone.localdate().strftime('%d/%m/%Y')
    if 'signature' in label or 'signed' in label:
        return '____________________________'
    if 'position' in label or 'capacity' in label:
        return 'Authorised Representative'
    return 'To be completed'


def fill_known_placeholders(text, bid_pack):
    if not bid_pack:
        return text
    company = bid_pack.company
    tender = bid_pack.tender
    replacements = {
        '[name of bidder]': company.name,
        '[name of tender]': tender.title,
        '[name of procuring entity]': tender.procuring_entity or '',
        '[insert date]': timezone.localdate().strftime('%d/%m/%Y'),
        '[insert name]': company.name,
        '[insert address]': company.address or '',
    }
    filled = str(text)
    for key, value in replacements.items():
        filled = re.sub(re.escape(key), value, filled, flags=re.I)
    return filled


def add_pdf_xml_preparation_checklist(elements, bid_document, style):
    elements.append(Paragraph('XML Requirement Checklist', style['Heading2']))
    response_items = bid_document.expected_response_lines
    if not response_items:
        response_items = ['Complete or attach the response required by this XML item.']

    rows = [['No.', 'What must be prepared / attached', 'TenderAI match', 'Action']]
    for index, item in enumerate(response_items, start=1):
        item_text = f'{bid_document.requirement} {item}'
        matched = match_company_document_for_text(bid_document.bid_pack, item_text)
        expected_type = expected_document_type_for_text(item_text)
        if matched:
            match_text = f'{matched.get_document_type_display()} - {matched.title}'
            action = 'Attached from company documents'
        elif looks_like_attachment_request(item):
            match_text = expected_type.label if expected_type else 'Not found'
            action = 'Upload this document type or attach manually'
        else:
            match_text = 'Manual response'
            action = 'Fill in / confirm'
        rows.append([
            str(index),
            Paragraph(item, style['BodyText']),
            Paragraph(match_text, style['BodyText']),
            Paragraph(action, style['BodyText']),
        ])
    elements.append(Table(rows, colWidths=[30, 210, 140, 95], style=table_style()))

    missing = [
        item for item in response_items
        if looks_like_attachment_request(item) and not match_company_document_for_text(bid_document.bid_pack, f'{bid_document.requirement} {item}')
    ]
    if missing:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph('Manual action still needed:', style['Heading3']))
        for item in missing:
            elements.append(Paragraph(f'- {item}', style['BodyText']))


def looks_like_attachment_request(text):
    return any(word in text.lower() for word in ['attach', 'certificate', 'evidence', 'documentary', 'compliance', 'printout', 'status'])


def expected_document_type_for_text(text):
    text = normalize_match_text(text)
    ncc_type = ncc_document_type_for_text(text)
    if ncc_type:
        return ncc_type
    for document_type, keywords in DOCUMENT_MATCH_RULES:
        if any(keyword in text for keyword in keywords):
            return document_type
    return None


def ncc_document_type_for_text(text):
    text = normalize_match_text(text)
    if 'national council' not in text and 'ncc' not in text:
        return None
    if re.search(r'\bcategory\s*(b|c)\b', text) or re.search(r'\bgrade\s*[1-4]\b', text):
        return CompanyDocument.DocumentType.NCC_B
    if re.search(r'\bncc\s*b\b', text) or 'ncc-b' in text:
        return CompanyDocument.DocumentType.NCC_B
    if re.search(r'\bncc\s*r\b', text) or 'ncc-r' in text or 'road' in text:
        return CompanyDocument.DocumentType.NCC_R
    if re.search(r'\bncc\s*e\b', text) or 'ncc-e' in text or 'electrical' in text:
        return CompanyDocument.DocumentType.NCC_E
    return None


def ncc_required_grade_for_text(text):
    text = normalize_match_text(text)
    if 'national council' not in text and 'ncc' not in text:
        return None
    match = re.search(r'\bgrade\s*(\d+)\s*(?:or\s*better|and\s*above|or\s*above)?', text)
    if match:
        return int(match.group(1))
    return None


def ncc_document_grade(document):
    if not document:
        return None
    text = document_match_text(document)
    match = re.search(r'\b(?:grade|g)\s*(\d+)\b', text)
    if match:
        return int(match.group(1))
    spaced_type_match = re.search(r'\bncc\s*[bre]?\s*(\d+)\b', text)
    if spaced_type_match:
        return int(spaced_type_match.group(1))
    return None


def ncc_grade_warning(bid_document):
    document = bid_document.matched_company_document
    required_grade = ncc_required_grade_for_text(f'{bid_document.requirement} {bid_document.expected_response}')
    actual_grade = ncc_document_grade(document)
    if required_grade is None or actual_grade is None:
        return ''
    if actual_grade > required_grade:
        return (
            f'{document.get_document_type_display()} Grade {actual_grade} is attached, '
            f'but the tender requires Grade {required_grade} or better. Upgrade/renew the NCC certificate before submission.'
        )
    return ''


def bid_document_attachment_warning(bid_document):
    document = bid_document.matched_company_document
    warnings = []
    if not document:
        return ''
    if document.is_expired:
        expiry = document.expiry_date.strftime('%d/%m/%Y') if document.expiry_date else 'the recorded expiry date'
        warnings.append(
            f'{document.get_document_type_display()} is attached but expired on {expiry}. Upload a current certificate before submission.'
        )
    grade_warning = ncc_grade_warning(bid_document)
    if grade_warning:
        warnings.append(grade_warning)
    return ' '.join(warnings)


def xml_required_document_evidence(tender):
    evidence = {}
    manual_items = []
    for item in expanded_tender_xml_items(tender):
        response_items = item.get('response_items') or []
        if not response_items and item.get('response'):
            response_items = [item.get('response')]
        if not response_items:
            response_items = [item.get('requirement') or item.get('title') or '']
        for response_item in response_items:
            text = f"{item.get('requirement', '')} {response_item}"
            if not looks_like_attachment_request(text):
                continue
            document_type = expected_document_type_for_text(text)
            if document_type:
                bucket = evidence.setdefault(document_type, {
                    'document_type': document_type,
                    'label': document_type.label,
                    'examples': [],
                    'count': 0,
                })
                bucket['count'] += 1
                if response_item and len(bucket['examples']) < 3:
                    bucket['examples'].append(response_item)
            elif response_item:
                manual_items.append(response_item)
    return {
        'typed': list(evidence.values()),
        'manual_items': manual_items[:8],
        'manual_count': len(manual_items),
    }


def document_gap_rows_for_company(tender, company):
    rows = []
    for item in xml_required_document_evidence(tender)['typed']:
        docs = list(company.documents.filter(document_type=item['document_type']).order_by('-expiry_date', '-uploaded_at'))
        active = [document for document in docs if not document.is_expired]
        expired = [document for document in docs if document.is_expired]
        if active:
            status = 'Ready'
            badge = 'success'
            action = 'Ready for attachment'
            document = active[0]
        elif expired:
            status = 'Expired'
            badge = 'danger'
            action = 'Renew or upload a current copy'
            document = expired[0]
        else:
            status = 'Missing'
            badge = 'warning'
            action = 'Upload this document'
            document = None
        rows.append({
            **item,
            'status': status,
            'badge': badge,
            'action': action,
            'document': document,
        })
    return rows


def add_pdf_placeholder_form(elements, bid_document, title, style):
    elements.append(Paragraph(title.upper(), style['FormTitle']))
    elements.append(Paragraph('This response is prepared in the format of the requirement captured from the solicitation document. Review and complete any blank fields before submission.', style['FormInstruction']))
    elements.append(Spacer(1, 6))
    bid_pack = bid_document.bid_pack
    rows = [
        ['Requirement', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Bidder response', Paragraph(default_response_for_requirement(bid_document), style['BodyText'])],
        ['Supporting document', Paragraph(supporting_document_label(bid_document), style['BodyText'])],
        ['Completion status', Paragraph('To be reviewed, signed, stamped, or attached as required by the solicitation document.', style['BodyText'])],
        ['Bidder', bid_pack.company.name],
        ['Authorised signature', '____________________________'],
        ['Date', timezone.localdate().strftime('%d/%m/%Y')],
    ]
    elements.append(Table(rows, colWidths=[145, 330], style=form_table_style()))
    if bid_document.expected_response_lines:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph('Items to confirm from XML / solicitation:', style['Heading3']))
        expected_rows = [['No.', 'Instruction / attachment required']]
        expected_rows.extend([
            [str(index), Paragraph(line, style['BodyText'])]
            for index, line in enumerate(bid_document.expected_response_lines, start=1)
        ])
        elements.append(Table(expected_rows, colWidths=[35, 440], style=form_table_style()))


def add_pdf_certificate_reference(elements, bid_document, style):
    document = bid_document.matched_company_document
    elements.append(Paragraph('DOCUMENTARY EVIDENCE / ATTACHMENT', style['FormTitle']))
    elements.append(Paragraph(
        f'The following company document is submitted in response to this requirement: {document.get_document_type_display()} - {document.title}.',
        style['BodyText'],
    ))
    rows = [
        ['Requirement', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Document type', document.get_document_type_display()],
        ['Title', document.title],
        ['Issue date', document.issue_date or '-'],
        ['Expiry date', document.expiry_date or '-'],
        ['Status', document.status_label],
        ['Bidder', bid_document.bid_pack.company.name],
    ]
    elements.append(Table(rows, colWidths=[145, 330], style=form_table_style()))


def add_pdf_qualification_form(elements, bid_pack, code, style, source_requirement=''):
    if code == 'EXP-2.4.1':
        add_pdf_general_experience_form(elements, bid_pack, style)
        return
    if code == 'MFR':
        add_pdf_manufacturer_authorisation_form(elements, bid_pack, style, source_requirement=source_requirement)
        return

    spec = build_qualification_form_spec(bid_pack, code)
    elements.append(Paragraph(spec['title'].upper(), style['FormTitle']))
    elements.append(Paragraph(f'Qualification factor: {spec.get("factor", "Qualification")}', style['FormInstruction']))
    if source_requirement:
        elements.append(Paragraph(f'Source requirement: {source_requirement}', style['FormInstruction']))
    elements.append(Spacer(1, 6))
    if 'headers' in spec:
        rows = [spec['headers'], *spec['rows']]
        col_width = 475 / len(spec['headers'])
        elements.append(Table(rows, colWidths=[col_width] * len(spec['headers']), style=form_table_style()))
    else:
        rows = [['Field', 'Bidder information / response']]
        rows.extend([[field, Paragraph(str(value), style['BodyText'])] for field, value in spec['rows']])
        elements.append(Table(rows, colWidths=[170, 305], style=form_table_style()))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('Declaration: The bidder confirms that the information provided in this form is true and complete, subject to final verification before submission.', style['BodyText']))
    elements.append(Table([
        ['Name of authorised representative', 'To be completed'],
        ['Signature and stamp', '____________________________'],
        ['Date', timezone.localdate().strftime('%d/%m/%Y')],
    ], colWidths=[190, 285], style=form_table_style()))


def add_pdf_delivery_warranty_form(elements, bid_document, style):
    bid_pack = bid_document.bid_pack
    tender = bid_pack.tender
    rows = [
        ['Tender', tender.title],
        ['Bidder', bid_pack.company.name],
        ['Requirement', Paragraph(bid_document.requirement, style['BodyText'])],
        ['Delivery / service period offered', 'To be completed in accordance with the solicitation schedule'],
        ['Warranty / undertaking offered', 'The bidder undertakes to comply with the warranty, delivery, and support requirements stated in the solicitation document.'],
        ['Exceptions', 'None, unless stated separately in the bid submission.'],
        ['Authorised signature', '____________________________'],
    ]
    elements.append(Paragraph('DELIVERY PERIOD / WARRANTY UNDERTAKING', style['FormTitle']))
    elements.append(Table(rows, colWidths=[165, 310], style=form_table_style()))


def add_pdf_technical_response_form(elements, bid_document, style):
    rows = [['No.', 'Technical requirement', 'Bidder response / offered specification']]
    response_items = bid_document.expected_response_lines or [bid_document.requirement]
    for index, item in enumerate(response_items, start=1):
        rows.append([
            str(index),
            Paragraph(item, style['BodyText']),
            Paragraph(prefilled_technical_response(item), style['BodyText']),
        ])
    elements.append(Paragraph('TECHNICAL RESPONSE SCHEDULE', style['FormTitle']))
    elements.append(Table(rows, colWidths=[35, 220, 220], style=form_table_style()))


def match_company_document_for_xml_item(bid_pack, item):
    text = ' '.join([
        str(item.get('requirement', '')),
        str(item.get('response', '')),
        ' '.join(item.get('response_items') or []),
    ]).lower()
    return match_company_document_for_text(bid_pack, text)


QUALIFICATION_FORM_ALIASES = {
    'eli 1.1': 'ELI 1.1',
    'eli-1.1': 'ELI 1.1',
    'bidder information sheet': 'ELI 1.1',
    'eli 1.2': 'ELI 1.2',
    'eli-1.2': 'ELI 1.2',
    'party to jv': 'ELI 1.2',
    'con-2': 'CON-2',
    'con 2': 'CON-2',
    'historical contract non-performance': 'CON-2',
    'con-3': 'CON-3',
    'con 3': 'CON-3',
    'current contract commitments': 'CCC',
    'form ccc': 'CCC',
    'fin-3.1': 'FIN-3.1',
    'fin 3.1': 'FIN-3.1',
    'financial situation': 'FIN-3.1',
    'fin-3.2': 'FIN-3.2',
    'fin3.2': 'FIN-3.2',
    'average annual turnover': 'FIN-3.2',
    'fin-3.3': 'FIN-3.3',
    'fin3.3': 'FIN-3.3',
    'financial resources': 'FIN-3.3',
    'exp-2.4.1': 'EXP-2.4.1',
    'general experience': 'EXP-2.4.1',
    'exp-2.4.2(a)': 'EXP-2.4.2(a)',
    'specific experience': 'EXP-2.4.2(a)',
    'exp-2.4.2(b)': 'EXP-2.4.2(b)',
    'key activities': 'EXP-2.4.2(b)',
    'per-1': 'PER-1',
    'proposed personnel': 'PER-1',
    'per-2': 'PER-2',
    'resume of proposed personnel': 'PER-2',
    'equipment': 'EQU',
    'manufacturer': 'MFR',
    'manufacturer authorisation': 'MFR',
    'manufacturer authorization': 'MFR',
}


def qualification_form_code_for_text(text):
    text = normalize_match_text(text)
    for alias, code in QUALIFICATION_FORM_ALIASES.items():
        if alias in text and code in QUALIFICATION_FORM_DEFINITIONS:
            return code
    return None


def supporting_document_label(bid_document):
    plan = experience_document_plan(f'{bid_document.requirement} {bid_document.expected_response}')
    if plan:
        return f'{plan["form_title"]}; attach {", ".join(plan["evidence"])}'
    document = bid_document.matched_company_document or match_company_document_for_text(
        bid_document.bid_pack,
        f'{bid_document.requirement} {bid_document.expected_response}',
    )
    if document:
        return f'{document.get_document_type_display()} - {document.title}'
    expected_type = expected_document_type_for_text(f'{bid_document.requirement} {bid_document.expected_response}')
    if expected_type:
        return f'{expected_type.label} to be attached'
    return 'Not applicable / manual confirmation'


def default_response_for_requirement(bid_document):
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    company = bid_document.bid_pack.company
    if 'bid validity' in text:
        return 'The bid validity period shall comply with the validity period stated in the solicitation document.'
    if 'payment terms' in text:
        return 'The bidder accepts the payment terms stated in the solicitation document, unless otherwise stated in the financial offer.'
    if 'price structure' in text:
        return 'The price structure shall be completed in the prescribed price schedule / commercial offer.'
    if 'availability' in text and 'station' in text:
        return 'To be completed with the bidder location / facility details applicable to this tender.'
    if looks_like_attachment_request(text):
        return f'{company.name} attaches or will attach the documentary evidence listed in this requirement.'
    return f'{company.name} confirms compliance with this requirement, subject to final review against the solicitation document.'


def prefilled_technical_response(item):
    text = normalize_match_text(item)
    if 'mandatory' in text:
        suffix = ' Mandatory requirement acknowledged.'
    else:
        suffix = ''
    if 'diesel' in text:
        return 'Compliant. Low sulphur diesel shall be supplied as specified.' + suffix
    if 'petrol' in text:
        return 'Compliant. Unleaded petrol shall be supplied as specified.' + suffix
    if 'specification' in text:
        return 'Compliant / to be completed with detailed offered specification.' + suffix
    if 'location' in text or 'availability' in text:
        return 'To be completed with location and availability evidence.' + suffix
    return 'Compliant / to be completed with bidder-specific details.' + suffix


DOCUMENT_MATCH_RULES = [
    (CompanyDocument.DocumentType.PACRA, [
        'pacra', 'certificate of incorporation', 'incorporation certificate',
        'schedule of directors', 'pacra printout', 'company registration',
    ]),
    (CompanyDocument.DocumentType.ZRA_TAX_CLEARANCE, [
        'tax clearance', 'zra tax', 'zra clearance', 'valid zra',
    ]),
    (CompanyDocument.DocumentType.TPIN_CERTIFICATE, ['tpin']),
    (CompanyDocument.DocumentType.NAPSA, ['napsa', 'pensions schemes authority']),
    (CompanyDocument.DocumentType.WORKERS_COMPENSATION, [
        'workers compensation', 'worker compensation', 'compensation fund',
        'workers compensation fund control board', 'wcfc',
    ]),
    (CompanyDocument.DocumentType.NCC_B, ['ncc b', 'ncc-b']),
    (CompanyDocument.DocumentType.NCC_R, ['ncc r', 'ncc-r']),
    (CompanyDocument.DocumentType.NCC_E, ['ncc e', 'ncc-e']),
    (CompanyDocument.DocumentType.NCC, ['ncc', 'national council for construction']),
    (CompanyDocument.DocumentType.ERB, [
        'erb', 'energy regulation board', 'erb rating', 'b+ or better',
    ]),
    (CompanyDocument.DocumentType.EIZ_CERTIFICATE, [
        'eiz', 'engineering institution of zambia', 'engineering institution',
    ]),
    (CompanyDocument.DocumentType.ROADWORTHINESS, [
        'roadworthiness', 'road worthy', 'vehicle compliance',
        'occupant safety', 'airbags', 'abs', 'seat belts', 'crash-test',
        'vehicle does not pose environmental',
    ]),
    (CompanyDocument.DocumentType.ZEMA_LICENSE, [
        'zema', 'zambia environmental management agency',
        'environmental compliance', 'environmental and social',
        'environmental certificate', 'environmental licence',
    ]),
    (CompanyDocument.DocumentType.ZPPA_REGISTRATION, [
        'zppa registration', 'public procurement registration',
    ]),
    (CompanyDocument.DocumentType.BANK_CONFIRMATION, [
        'bank confirmation', 'bank letter', 'financial resources',
        'liquid assets', 'credit line', 'line of credit', 'bank statement',
    ]),
    (CompanyDocument.DocumentType.AUDITED_FINANCIALS, [
        'audited financial', 'financial statements', 'average annual turnover',
        'annual turnover', 'balance sheet', 'income statement',
    ]),
    (CompanyDocument.DocumentType.DELIVERY_EVIDENCE, [
        'delivery notes', 'delivery note', 'lpo', 'local purchase order',
        'grn', 'goods received note', 'award letters', 'award letter',
        'executed similar assignments', 'government of zambia',
    ]),
    (CompanyDocument.DocumentType.TRAINING_PROGRAMME, [
        'training programme', 'training program', 'technical capacity building',
        'capacity building', 'machine orientation', 'operator orientation',
    ]),
    (CompanyDocument.DocumentType.WARRANTY_UNDERTAKING, [
        'warranty undertaking', 'warranty letter', 'warranty period',
    ]),
    (CompanyDocument.DocumentType.PAST_CONTRACT, [
        'similar experience', 'past contract', 'traceable references',
        'reference letters', 'specific experience', 'general experience',
        'contracts that have been successfully executed',
    ]),
    (CompanyDocument.DocumentType.COMPANY_PROFILE, ['company profile']),
]


def normalize_match_text(value):
    return re.sub(r'\s+', ' ', str(value or '').replace('_', ' ').lower()).strip()


def document_match_text(document):
    filename = document.file.name if document.file else ''
    return normalize_match_text(' '.join([
        document.document_type,
        document.get_document_type_display(),
        document.title,
        document.notes,
        filename,
    ]))


def rank_document(document, score):
    expiry_rank = 0 if not document.is_expired else -3
    has_file_rank = 1 if document.file else 0
    expiry = document.expiry_date or document.issue_date
    date_rank = expiry.toordinal() if expiry else 0
    upload_rank = document.uploaded_at.timestamp() if document.uploaded_at else 0
    return (score + expiry_rank + has_file_rank, date_rank, upload_rank)


def best_document_for_type(company, document_type):
    docs = list(company.documents.filter(document_type=document_type))
    if not docs:
        return None
    return max(docs, key=lambda document: rank_document(document, 10))


def best_document_by_keywords(company, keywords):
    candidates = []
    for document in company.documents.all():
        haystack = document_match_text(document)
        score = sum(2 for keyword in keywords if keyword in haystack)
        if document.document_type == CompanyDocument.DocumentType.OTHER and score:
            score += 1
        if score:
            candidates.append((rank_document(document, score), document))
    if not candidates:
        return None
    return max(candidates, key=lambda item: item[0])[1]


def match_company_document_for_text(bid_pack, text):
    text = normalize_match_text(text)
    ncc_type = ncc_document_type_for_text(text)
    if ncc_type:
        return best_document_for_type(bid_pack.company, ncc_type)
    for document_type, keywords in DOCUMENT_MATCH_RULES:
        if any(keyword in text for keyword in keywords):
            exact_match = best_document_for_type(bid_pack.company, document_type)
            if exact_match:
                return exact_match
            fallback_match = best_document_by_keywords(bid_pack.company, keywords)
            if fallback_match:
                return fallback_match
    return None


def append_matched_certificate_pdf(main_pdf, bid_document):
    document = bid_document.matched_company_document
    if not document or not document.file or not document.file.name.lower().endswith('.pdf'):
        return main_pdf
    writer = PdfWriter()
    append_pdf_bytes(writer, main_pdf)
    try:
        with document.file.open('rb') as handle:
            reader = PdfReader(handle)
            for page in reader.pages:
                writer.add_page(page)
    except Exception:
        return main_pdf
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def generated_bid_documents_index_pdf(bid_pack):
    style = styles()
    elements = []
    elements.append(Paragraph('XML Bid Pack Assembly Instructions', style['Heading1']))
    elements.append(Paragraph(
        'This pack follows the Tender Structure XML order only. Each item starts with an instruction page explaining what is required, followed by the generated form, letter, table, missing-document notice, or attached certificate.',
        style['BodyText'],
    ))
    elements.append(Spacer(1, 8))
    rows = [['No.', 'Part', 'Requirement', 'What TenderAI prepared']]
    for bid_document in bid_pack.bid_documents.order_by('order'):
        rows.append([
            str(bid_document.order),
            Paragraph(bid_document.envelope or '-', style['BodyText']),
            Paragraph(bid_document.requirement, style['BodyText']),
            Paragraph(bid_document_action_label(bid_document), style['BodyText']),
        ])
    elements.append(Table(rows, colWidths=[34, 120, 210, 111], style=form_table_style()))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
        'If a row says a required certificate is missing, upload that company document and regenerate these XML bid documents before submission.',
        style['FormInstruction'],
    ))
    return build_pdf_response(elements, f'Generated Bid Documents - {bid_pack.tender.title}')


def generated_envelope_divider_pdf(bid_pack, envelope, part_number):
    style = styles()
    divider_title = style['Title'].clone(f'EnvelopeDivider{part_number}')
    divider_title.alignment = 1
    divider_title.fontSize = 22
    divider_title.leading = 28
    divider_title.spaceAfter = 12
    body = style['BodyText'].clone(f'EnvelopeDividerBody{part_number}')
    body.alignment = 1
    elements = [
        Spacer(1, A4[1] * 0.27),
        Paragraph(f'PART {part_number}', style['Heading1']),
        Paragraph(str(envelope or 'Bid Requirements').upper(), divider_title),
        Paragraph('Documents in this section follow the matching XML envelope / solicitation schedule.', body),
        Spacer(1, 18),
        Paragraph(f'Bidder: {bid_pack.company.name}', body),
        Paragraph(f'Tender: {bid_pack.tender.title}', body),
    ]
    return build_pdf_response(elements, f'{envelope} - Divider')


def prepared_output_for_bid_document(bid_document):
    text = normalize_match_text(f'{bid_document.requirement} {bid_document.expected_response}')
    form_code = qualification_form_code_for_text(text)
    if 'bid declaration' in text or 'bid securing' in text or 'bid security' in text:
        return 'Bid-Securing Declaration form'
    if is_letter_of_bid_requirement(text):
        if is_price_schedule_requirement(text):
            return 'Letter of Bid plus Price Schedule'
        return 'Letter of Bid / Tender Submission Letter'
    if is_price_schedule_requirement(text):
        return 'Signed quotation / price schedule'
    if is_power_of_attorney_requirement(text):
        return 'Power of Attorney / Signatory Authorisation form'
    if form_code:
        if form_code == 'MFR':
            return "Manufacturer's Authorisation form"
        return f'Qualification form {form_code}'
    if bid_document.matched_company_document:
        return f'Certificate attachment: {bid_document.matched_company_document.get_document_type_display()}'
    if any(word in text for word in ['technical', 'specification', 'capacity', 'compliance']):
        return 'Technical response schedule'
    if any(word in text for word in ['commercial', 'price', 'payment', 'delivery', 'warranty']):
        return 'Commercial / undertaking schedule'
    return 'Prepared response form'


def generate_combined_bid_documents_pdf(bid_pack):
    if not bid_pack.bid_documents.exists():
        generate_xml_bid_documents(bid_pack)
    writer = PdfWriter()
    append_pdf_bytes(writer, generated_bid_documents_index_pdf(bid_pack))
    current_envelope = None
    part_number = 0
    for bid_document in bid_pack.bid_documents.order_by('order'):
        if not bid_document.generated_pdf:
            continue
        try:
            if bid_document.envelope != current_envelope:
                current_envelope = bid_document.envelope
                part_number += 1
                append_pdf_bytes(writer, generated_envelope_divider_pdf(bid_pack, current_envelope, part_number))
            with bid_document.generated_pdf.open('rb') as handle:
                reader = PdfReader(handle)
                for page in reader.pages:
                    writer.add_page(page)
        except Exception:
            continue
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def safe_slug(value):
    text = re.sub(r'[^a-zA-Z0-9]+', '-', str(value)).strip('-').lower()
    return text[:70] or 'document'


def table_style():
    return TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#d9e2ec')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#edf4f3')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('TOPPADDING', (0, 0), (-1, -1), 7),
    ])


def form_table_style():
    return TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.55, colors.HexColor('#333333')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
        ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEADING', (0, 0), (-1, -1), 11),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ])


def plain_field_table_style():
    return TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTNAME', (0, 0), (0, -1), 'Times-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('LEADING', (0, 0), (-1, -1), 13),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ])


def price_schedule_outer_style():
    return TableStyle([
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('LINEBEFORE', (1, 0), (1, 0), 0.8, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ])


def price_schedule_table_style():
    return TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.65, colors.black),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTNAME', (0, 1), (-1, 1), 'Times-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 7),
        ('LEADING', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ])


def price_schedule_total_style():
    return TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ])
